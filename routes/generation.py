import asyncio
import json
import logging
import re
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

from fastapi import APIRouter, Request
from fastapi.responses import JSONResponse, Response, StreamingResponse
import httpx

import services.config as _svc_config
from services.config import FALLBACK_MODEL, OLLAMA, PROJECTS_DIR
from services.system import _sse_event, _PDF_CSS
from services.queue import is_busy, operation_busy_response, queued_sse_stream
from services.projects import (
    _resolve_project_path,
    _get_project_knowledge_context,
    _parse_code_files,
    _slugify,
    _parse_workflow_plan,
)
from services.research import (
    _generate_search_queries,
    _do_web_search,
    _fetch_page_content,
    _extract_findings,
    _synthesize_summary,
    _save_markdown_as_pdf,
)
from services.media import (
    _MARKDOWN_AVAILABLE,
    _WEASYPRINT_AVAILABLE,
    _DOCX_AVAILABLE,
)

# Optional module imports for export handlers

try:
    import markdown as _markdown_lib
except ImportError:
    _markdown_lib = None

try:
    import weasyprint as _weasyprint
except ImportError:
    _weasyprint = None

try:
    import docx as _docx_module
except ImportError:
    _docx_module = None

log = logging.getLogger("offlineai")

router = APIRouter()


# ── Two-phase research pipeline ──────────────────────────────────────

@router.post("/api/projects/{project_id}/research/plan")
async def research_plan(project_id: str, request: Request):
    """Generate search queries for a research topic WITHOUT executing them.
    
    Returns a list of queries that the user can review, edit, and toggle
    before executing via /research/execute.
    """
    import uuid

    body = await request.json()
    topic = (body.get("topic") or "").strip()
    if not topic:
        return JSONResponse({"error": "No topic provided"}, status_code=400)

    depth = body.get("depth", "standard")
    model = body.get("model", _svc_config.FALLBACK_MODEL)
    num_queries = {"quick": 3, "standard": 5, "deep": 8}.get(depth, 5)

    project_path = _resolve_project_path(project_id)
    if project_path is None or not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    try:
        queries = await _generate_search_queries(topic, num_queries, model)
    except Exception:
        queries = [topic]

    return {
        "topic": topic,
        "depth": depth,
        "queries": [
            {"id": str(uuid.uuid4())[:8], "text": q, "enabled": True}
            for q in queries
        ],
        "estimated_sources": len(queries) * 2,  # ~2 sources per query
    }


@router.post("/api/projects/{project_id}/research/execute")
async def research_execute(project_id: str, request: Request):
    """Execute research with user-approved queries. Streams SSE progress.
    
    Accepts the queries array from /research/plan (possibly edited/toggled).
    Only executes queries where enabled=True.
    """
    body = await request.json()
    topic = (body.get("topic") or "").strip()
    if not topic:
        return JSONResponse({"error": "No topic provided"}, status_code=400)

    queries_list = body.get("queries", [])
    enabled_queries = [q["text"] for q in queries_list if q.get("enabled", True) and q.get("text", "").strip()]
    if not enabled_queries:
        return JSONResponse({"error": "No enabled queries to execute"}, status_code=400)

    depth = body.get("depth", "standard")
    model = body.get("model", _svc_config.FALLBACK_MODEL)

    # Configurable research parameters
    max_results_per_query = min(int(body.get("max_results_per_query", 2)), 5)
    max_page_chars = min(int(body.get("max_page_chars", 5000)), 8000)
    excluded_domains = body.get("excluded_domains", [])
    if not isinstance(excluded_domains, list):
        excluded_domains = []

    project_path = _resolve_project_path(project_id)
    if project_path is None or not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    knowledge_file = project_path / "knowledge.json"
    if not knowledge_file.exists():
        return JSONResponse({"error": "Invalid project"}, status_code=404)

    async def execute_stream():
        try:
            yield _sse_event({"type": "status", "message": f"Executing research: {topic} ({len(enabled_queries)} queries)"})
            yield _sse_event({"type": "phase", "phase": "searching"})
            log.info("Research execute: %s with %d queries", topic, len(enabled_queries))

            all_sources = []
            all_page_content = []

            for i, query in enumerate(enabled_queries):
                yield _sse_event({"type": "status", "message": f"Searching ({i+1}/{len(enabled_queries)}): {query}"})

                try:
                    search_results = await _do_web_search(query, max_results=5)
                except Exception:
                    search_results = []
                    yield _sse_event({"type": "status", "message": f"⚠ Search failed for: {query}"})

                # Filter excluded domains
                if excluded_domains:
                    search_results = [
                        r for r in search_results
                        if not any(domain in r.get("href", "") for domain in excluded_domains)
                    ]

                yield _sse_event({"type": "search", "query": query, "results_count": len(search_results)})

                for result in search_results[:max_results_per_query]:
                    url = result.get("href", "")
                    title = result.get("title", "")
                    snippet = result.get("body", "")
                    all_sources.append({
                        "url": url, "title": title, "snippet": snippet,
                        "fetched_at": datetime.now(timezone.utc).isoformat(),
                    })
                    yield _sse_event({"type": "source", "url": url, "title": title, "snippet": snippet})

            # Phase: reading pages
            yield _sse_event({"type": "phase", "phase": "reading"})
            yield _sse_event({"type": "status", "message": f"Reading {len(all_sources)} sources..."})

            for source in all_sources:
                page = await _fetch_page_content(source["url"], max_chars=max_page_chars)
                if page:
                    all_page_content.append(page)

            # Phase: extracting findings
            yield _sse_event({"type": "phase", "phase": "extracting"})
            if all_page_content:
                yield _sse_event({"type": "status", "message": f"Extracting findings from {len(all_page_content)} pages..."})
                findings_text = await _extract_findings(topic, all_page_content, model)
                yield _sse_event({"type": "finding", "text": findings_text})
            else:
                findings_text = "No source content available."
                yield _sse_event({"type": "status", "message": "No pages could be fetched"})

            # Phase: synthesizing
            yield _sse_event({"type": "phase", "phase": "synthesizing"})
            yield _sse_event({"type": "status", "message": "Writing research summary..."})
            summary = await _synthesize_summary(topic, findings_text, all_sources, model)

            # Emit source_map for citation rendering
            for i, src in enumerate(all_sources[:15], 1):
                yield _sse_event({"type": "source_map", "index": i, "title": src.get("title", ""), "url": src.get("url", "")})

            # Save files
            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            topic_slug = _slugify(topic)[:40]
            notes_dir = project_path / "notes"
            notes_dir.mkdir(exist_ok=True)
            summary_file = notes_dir / f"{timestamp}-{topic_slug}.md"
            summary_file.write_text(summary, encoding="utf-8")

            pdf_path = notes_dir / f"{timestamp}-{topic_slug}.pdf"
            try:
                await asyncio.to_thread(_save_markdown_as_pdf, summary, pdf_path, topic)
            except Exception:
                pass

            # Update knowledge.json
            knowledge = json.loads(knowledge_file.read_text(encoding="utf-8"))
            knowledge.setdefault("findings", []).append({
                "topic": topic,
                "summary": findings_text[:2000],
                "sources": [s.get("url", "") for s in all_sources[:10]],
                "timestamp": datetime.now(timezone.utc).isoformat(),
            })
            for s in all_sources:
                knowledge.setdefault("sources", []).append({
                    "title": s.get("title", ""),
                    "url": s.get("url", ""),
                    "timestamp": s.get("fetched_at", ""),
                })
            knowledge_file.write_text(json.dumps(knowledge, indent=2, ensure_ascii=False), encoding="utf-8")

            # Reindex for FTS5
            try:
                from services.knowledge_store import reindex_project
                reindex_project(project_id)
            except Exception:
                pass

            yield _sse_event({"type": "phase", "phase": "done"})
            yield _sse_event({
                "type": "done",
                "message": f"Research complete: {len(all_sources)} sources, {len(all_page_content)} pages read",
                "summary_file": f"notes/{summary_file.name}",
            })

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        queued_sse_stream(execute_stream()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── Legacy single-step research (backward compatibility) ─────────────

@router.post("/api/projects/{project_id}/research")
async def research_project(project_id: str, request: Request):
    """Autonomous multi-step research agent. Streams SSE progress."""
    body = await request.json()
    topic = (body.get("topic") or "").strip()
    if not topic:
        return JSONResponse({"error": "No topic provided"}, status_code=400)

    depth = body.get("depth", "standard")
    model = body.get("model", _svc_config.FALLBACK_MODEL)
    num_queries = {"quick": 3, "standard": 5, "deep": 8}.get(depth, 5)

    # Configurable research parameters
    max_results_per_query = min(int(body.get("max_results_per_query", 2)), 5)
    max_page_chars = min(int(body.get("max_page_chars", 5000)), 8000)
    excluded_domains = body.get("excluded_domains", [])
    if not isinstance(excluded_domains, list):
        excluded_domains = []

    project_path = _resolve_project_path(project_id)
    if project_path is None or not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    knowledge_file = project_path / "knowledge.json"
    if not knowledge_file.exists():
        return JSONResponse({"error": "Invalid project"}, status_code=404)

    async def research_stream():
        try:
            yield _sse_event({"type": "status", "message": f"Planning research on: {topic}"})
            log.info("Research started: %s depth=%s", topic, depth)

            try:
                queries = await _generate_search_queries(topic, num_queries, model)
            except Exception:
                queries = [topic]
            yield _sse_event({"type": "status", "message": f"Generated {len(queries)} search queries"})

            all_sources = []
            all_page_content = []
            search_errors = 0

            for i, query in enumerate(queries):
                yield _sse_event({"type": "status", "message": f"Searching ({i+1}/{len(queries)}): {query}"})

                try:
                    search_results = await _do_web_search(query, max_results=max_results_per_query)
                except Exception:
                    search_results = []
                    search_errors += 1
                    yield _sse_event({"type": "status", "message": f"⚠ Search failed for: {query}, continuing..."})

                yield _sse_event({"type": "search", "query": query, "results_count": len(search_results)})

                # Filter excluded domains
                if excluded_domains:
                    search_results = [
                        r for r in search_results
                        if not any(domain in r.get("href", "") for domain in excluded_domains)
                    ]

                for result in search_results[:max_results_per_query]:
                    url = result.get("href", "")
                    title = result.get("title", "")
                    snippet = result.get("body", "")

                    all_sources.append({
                        "url": url,
                        "title": title,
                        "snippet": snippet,
                        "fetched_at": datetime.now(timezone.utc).isoformat(),
                    })

                    try:
                        page_text = await _fetch_page_content(url, max_chars=max_page_chars)
                        if page_text:
                            all_page_content.append(f"Source: {title} ({url})\n{page_text[:max_page_chars]}")
                            yield _sse_event({"type": "source", "message": f"Read: {title}"})
                    except Exception:
                        yield _sse_event({"type": "status", "message": f"⚠ Could not read: {title[:50]}"})

            if not all_sources and not all_page_content:
                yield _sse_event({"type": "error", "error": f"All {len(queries)} searches failed. Check your internet connection."})
                return

            if search_errors > 0:
                yield _sse_event({"type": "status", "message": f"Completed with {search_errors} failed search(es), {len(all_sources)} sources found"})

            yield _sse_event({"type": "status", "message": "Analyzing sources and extracting findings..."})

            try:
                findings_text = await _extract_findings(topic, all_page_content, model)
            except Exception as exc:
                findings_text = "Key findings from search results:\n\n" + "\n\n".join(
                    f"- {s.get('title', 'Unknown')}: {s.get('snippet', '')}" for s in all_sources[:10]
                )
                yield _sse_event({"type": "status", "message": f"⚠ Analysis partially failed ({exc}), using raw findings"})

            yield _sse_event({"type": "finding", "text": findings_text[:2000]})

            yield _sse_event({"type": "status", "message": "Writing comprehensive summary..."})

            # Emit numbered source map so frontend can render clickable citation links
            for i, src in enumerate(all_sources[:15], 1):
                yield _sse_event({"type": "source_map", "index": i, "title": src.get("title", "Unknown"), "url": src.get("url", "")})

            try:
                summary = await _synthesize_summary(topic, findings_text, all_sources, model)
            except Exception as exc:
                summary = f"# Research: {topic}\n\n## Findings\n\n{findings_text}\n\n## Sources\n\n" + "\n".join(
                    f"- [{s.get('title', 'Unknown')}]({s.get('url', '')})" for s in all_sources[:10]
                )
                yield _sse_event({"type": "status", "message": f"⚠ Summary generation failed ({exc}), using raw findings"})

            yield _sse_event({"type": "status", "message": "Saving research to project..."})

            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            topic_slug = _slugify(topic)[:40]
            note_filename = f"{timestamp}-{topic_slug}.md"
            notes_dir = project_path / "notes"
            notes_dir.mkdir(exist_ok=True)
            note_path = notes_dir / note_filename
            note_path.write_text(summary, encoding="utf-8")

            pdf_filename = f"{timestamp}-{topic_slug}.pdf"
            try:
                await asyncio.to_thread(_save_markdown_as_pdf, summary, notes_dir / pdf_filename, topic)
            except Exception:
                pass

            knowledge = json.loads(knowledge_file.read_text(encoding="utf-8"))

            existing_urls = {s["url"] for s in knowledge.get("sources", [])}
            for src in all_sources:
                if src["url"] not in existing_urls:
                    knowledge.setdefault("sources", []).append(src)
                    existing_urls.add(src["url"])

            knowledge.setdefault("findings", []).append({
                "topic": topic,
                "summary": findings_text[:2000],
                "sources": [s["url"] for s in all_sources[:10]],
                "timestamp": datetime.now(timezone.utc).isoformat(),
            })

            knowledge_file.write_text(json.dumps(knowledge, indent=2, ensure_ascii=False), encoding="utf-8")

            log.info("Research complete: %s — %d sources", topic, len(all_sources))
            yield _sse_event({"type": "content", "text": summary})

            yield _sse_event({
                "type": "done",
                "message": "Research complete!",
                "summary_file": f"notes/{note_filename}",
            })

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        queued_sse_stream(research_stream()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.post("/api/projects/{project_id}/generate-document")
async def generate_document(project_id: str, request: Request):
    """Generate a Markdown document, optionally using project knowledge. Streams SSE."""
    body = await request.json()
    topic = (body.get("topic") or "").strip()
    if not topic:
        return JSONResponse({"error": "No topic provided"}, status_code=400)

    project_path = _resolve_project_path(project_id)
    if project_path is None or not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    doc_type = body.get("type", "report")
    model = body.get("model", _svc_config.FALLBACK_MODEL)
    use_knowledge = body.get("use_knowledge", True)

    async def doc_stream():
        try:
            yield _sse_event({"type": "status", "message": f"Generating {doc_type}: {topic}"})

            context = ""
            if use_knowledge:
                context = _get_project_knowledge_context(project_id, max_chars=6000)

            prompt = f"""Write a comprehensive {doc_type} about: "{topic}"

{f'Use this research context to inform your writing:{chr(10)}{context}' if context else ''}

Format as a well-structured Markdown document with:
- A clear title (# heading)
- An introduction/executive summary
- Organized sections with ## subheadings
- Bullet points and numbered lists where appropriate
- A conclusion or summary section
- A "## References" section at the end listing all cited sources as a numbered list

If sources are provided above, use inline citation markers [1], [2], etc. to reference them. Place the citation immediately after the claim it supports. In the References section, list each cited source as: [N] Title — URL. Only cite sources from the provided list.

Be thorough, detailed, and well-organized."""

            yield _sse_event({"type": "status", "message": "Writing document..."})

            content = ""
            async with httpx.AsyncClient(timeout=httpx.Timeout(connect=5.0, read=None, write=120.0, pool=5.0)) as client:
                async with client.stream("POST", f"{_svc_config.OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": True,
                    "options": {"temperature": 0.5, "num_predict": 8192, "num_ctx": 32768},
                }) as resp:
                    async for line in resp.aiter_lines():
                        line = line.strip()
                        if not line:
                            continue
                        try:
                            data = json.loads(line)
                            token = data.get("message", {}).get("content", "")
                            if token:
                                content += token
                                yield _sse_event({"type": "token", "text": token})
                            if data.get("done"):
                                break
                        except json.JSONDecodeError:
                            continue

            if not content:
                yield _sse_event({"type": "error", "error": "Model returned empty content"})
                return

            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            topic_slug = _slugify(topic)[:40]
            filename = f"{timestamp}-{topic_slug}.md"
            output_dir = project_path / "output"
            output_dir.mkdir(exist_ok=True)
            file_path = output_dir / filename
            file_path.write_text(content, encoding="utf-8")
            log.info("Document generated: %s", filename)

            pdf_filename = f"{timestamp}-{topic_slug}.pdf"
            pdf_path = output_dir / pdf_filename
            try:
                await asyncio.to_thread(_save_markdown_as_pdf, content, pdf_path, topic)
                yield _sse_event({"type": "status", "message": f"PDF saved: output/{pdf_filename}"})
            except Exception:
                pass

            yield _sse_event({"type": "content", "text": content})
            yield _sse_event({"type": "done", "message": "Document generated", "file_path": f"output/{filename}"})

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        queued_sse_stream(doc_stream()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.post("/api/projects/{project_id}/export-pdf")
async def export_pdf(project_id: str, request: Request):
    """Convert a Markdown file to a human-readable PDF."""
    body = await request.json()
    file_path = (body.get("file_path") or "").strip()
    if not file_path:
        return JSONResponse({"error": "No file_path provided"}, status_code=400)

    resolved = _resolve_project_path(project_id, file_path)
    if resolved is None or not resolved.exists():
        return JSONResponse({"error": "File not found"}, status_code=404)

    md_content = resolved.read_text(encoding="utf-8")

    if _MARKDOWN_AVAILABLE:
        html_body = _markdown_lib.markdown(md_content, extensions=["tables", "fenced_code", "toc", "nl2br"])
    else:
        html_body = f"<pre>{md_content}</pre>"

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{Path(file_path).stem}</title>
<style>
{_PDF_CSS}
</style></head><body>{html_body}</body></html>"""

    if _WEASYPRINT_AVAILABLE:
        def _generate_pdf():
            doc = _weasyprint.HTML(string=html)
            return doc.write_pdf()

        pdf_bytes = await asyncio.to_thread(_generate_pdf)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{Path(file_path).stem}.pdf"'},
        )

    return Response(
        content=html.encode("utf-8"),
        media_type="text/html",
        headers={"Content-Disposition": f'attachment; filename="{Path(file_path).stem}.html"'},
    )


@router.post("/api/projects/{project_id}/export-docx")
async def export_docx(project_id: str, request: Request):
    """Convert a Markdown file to a DOCX document."""
    body = await request.json()
    file_path = (body.get("file_path") or "").strip()
    if not file_path:
        return JSONResponse({"error": "No file_path provided"}, status_code=400)

    resolved = _resolve_project_path(project_id, file_path)
    if resolved is None or not resolved.exists():
        return JSONResponse({"error": "File not found"}, status_code=404)

    if not _DOCX_AVAILABLE:
        return JSONResponse({"error": "python-docx is not installed"}, status_code=500)

    md_content = resolved.read_text(encoding="utf-8")

    def _build_docx(md_text: str) -> bytes:
        doc = _docx_module.Document()

        in_code_block = False
        code_lines: list[str] = []

        for line in md_text.split("\n"):
            # --- fenced code block toggle ---
            if line.strip().startswith("```"):
                if in_code_block:
                    # close code block
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = _docx_module.shared.Pt(9)
                    p.style = doc.styles["Normal"]
                    fmt = p.paragraph_format
                    fmt.space_before = _docx_module.shared.Pt(4)
                    fmt.space_after = _docx_module.shared.Pt(4)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            stripped = line.strip()

            # --- blank lines ---
            if not stripped:
                continue

            # --- headings ---
            heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
            if heading_match:
                level = min(len(heading_match.group(1)), 3)
                doc.add_heading(heading_match.group(2).strip(), level=level)
                continue

            # --- bullet lists ---
            bullet_match = re.match(r"^[\-\*]\s+(.*)", stripped)
            if bullet_match:
                doc.add_paragraph(bullet_match.group(1), style="List Bullet")
                continue

            # --- numbered lists ---
            num_match = re.match(r"^\d+[\.\)]\s+(.*)", stripped)
            if num_match:
                doc.add_paragraph(num_match.group(1), style="List Number")
                continue

            # --- horizontal rule ---
            if re.match(r"^[-*_]{3,}\s*$", stripped):
                doc.add_paragraph("─" * 40)
                continue

            # --- blockquote ---
            if stripped.startswith(">"):
                text = re.sub(r"^>\s?", "", stripped)
                p = doc.add_paragraph(text)
                fmt = p.paragraph_format
                fmt.left_indent = _docx_module.shared.Inches(0.5)
                continue

            # --- normal paragraph ---
            doc.add_paragraph(stripped)

        # flush any unclosed code block
        if code_lines:
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = _docx_module.shared.Pt(9)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    docx_bytes = await asyncio.to_thread(_build_docx, md_content)
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{Path(file_path).stem}.docx"'},
    )


@router.post("/api/projects/{project_id}/export-html")
async def export_html(project_id: str, request: Request):
    """Convert a Markdown file to a standalone styled HTML document."""
    body = await request.json()
    file_path = (body.get("file_path") or "").strip()
    if not file_path:
        return JSONResponse({"error": "No file_path provided"}, status_code=400)

    resolved = _resolve_project_path(project_id, file_path)
    if resolved is None or not resolved.exists():
        return JSONResponse({"error": "File not found"}, status_code=404)

    md_content = resolved.read_text(encoding="utf-8")

    if _MARKDOWN_AVAILABLE:
        html_body = _markdown_lib.markdown(md_content, extensions=["tables", "fenced_code", "toc", "nl2br"])
    else:
        html_body = f"<pre>{md_content}</pre>"

    html = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>{Path(file_path).stem}</title>
<style>
body {{ max-width: 860px; margin: 40px auto; padding: 0 24px; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Helvetica Neue', sans-serif; font-size: 16px; line-height: 1.7; color: #1a1a1a; background: #fff; }}
h1 {{ font-size: 2em; border-bottom: 2px solid #333; padding-bottom: 8px; margin-top: 0; margin-bottom: 16px; }}
h2 {{ font-size: 1.5em; border-bottom: 1px solid #ccc; padding-bottom: 6px; margin-top: 32px; color: #2c3e50; }}
h3 {{ font-size: 1.17em; margin-top: 24px; color: #34495e; }}
p {{ margin: 10px 0; }}
ul, ol {{ margin: 10px 0; padding-left: 28px; }}
li {{ margin: 4px 0; }}
code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; font-family: 'SF Mono', 'Fira Code', Menlo, monospace; }}
pre {{ background: #f8f8f8; padding: 16px 20px; border-radius: 6px; overflow-x: auto; border: 1px solid #e8e8e8; font-size: 0.88em; line-height: 1.5; }}
pre code {{ background: none; padding: 0; }}
table {{ border-collapse: collapse; width: 100%; margin: 16px 0; font-size: 0.95em; }}
th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
th {{ background: #f0f0f0; font-weight: 600; }}
tr:nth-child(even) {{ background: #fafafa; }}
blockquote {{ border-left: 4px solid #3498db; margin: 16px 0; padding: 10px 20px; color: #555; background: #f9fbfd; border-radius: 0 4px 4px 0; }}
a {{ color: #2980b9; text-decoration: none; }}
a:hover {{ text-decoration: underline; }}
hr {{ border: none; border-top: 1px solid #ddd; margin: 24px 0; }}
img {{ max-width: 100%; height: auto; }}
@media (prefers-color-scheme: dark) {{
  body {{ background: #1a1a2e; color: #e0e0e0; }}
  h1 {{ border-bottom-color: #555; }}
  h2 {{ border-bottom-color: #444; color: #8ab4f8; }}
  h3 {{ color: #8ab4f8; }}
  code {{ background: #2a2a3e; }}
  pre {{ background: #252540; border-color: #333; }}
  th {{ background: #2a2a3e; }}
  tr:nth-child(even) {{ background: #1f1f35; }}
  td, th {{ border-color: #333; }}
  blockquote {{ background: #1f1f35; border-left-color: #3498db; color: #aaa; }}
  a {{ color: #5dade2; }}
}}
</style></head><body>{html_body}</body></html>"""

    return Response(
        content=html.encode("utf-8"),
        media_type="text/html",
        headers={"Content-Disposition": f'attachment; filename="{Path(file_path).stem}.html"'},
    )


@router.post("/api/projects/{project_id}/generate-data")
async def generate_data(project_id: str, request: Request):
    """Generate structured data (CSV or JSON). Streams SSE."""
    body = await request.json()
    topic = (body.get("topic") or "").strip()
    if not topic:
        return JSONResponse({"error": "No topic provided"}, status_code=400)

    project_path = _resolve_project_path(project_id)
    if project_path is None or not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    data_format = body.get("format", "csv")
    model = body.get("model", _svc_config.FALLBACK_MODEL)

    async def data_stream():
        try:
            yield _sse_event({"type": "status", "message": f"Generating {data_format.upper()} data: {topic}"})

            if data_format == "csv":
                prompt = f"""Generate a CSV dataset about: "{topic}"

Rules:
- First row must be column headers
- Use commas as delimiters
- Wrap fields containing commas in double quotes
- Include at least 10 rows of meaningful data
- Return ONLY the CSV content, no explanations or markdown formatting"""
            else:
                prompt = f"""Generate a JSON dataset about: "{topic}"

Rules:
- Return a JSON array of objects
- Each object should have consistent keys
- Include at least 10 items with meaningful data
- Return ONLY valid JSON, no explanations or markdown formatting"""

            content = ""
            async with httpx.AsyncClient(timeout=httpx.Timeout(connect=5.0, read=None, write=120.0, pool=5.0)) as client:
                async with client.stream("POST", f"{_svc_config.OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": True,
                    "options": {"temperature": 0.3, "num_predict": 4096, "num_ctx": 32768},
                }) as resp:
                    async for line in resp.aiter_lines():
                        line = line.strip()
                        if not line:
                            continue
                        try:
                            data = json.loads(line)
                            token = data.get("message", {}).get("content", "")
                            if token:
                                content += token
                                yield _sse_event({"type": "token", "text": token})
                            if data.get("done"):
                                break
                        except json.JSONDecodeError:
                            continue

            if not content:
                yield _sse_event({"type": "error", "error": "Model returned empty content"})
                return

            content = content.strip()
            if content.startswith("```"):
                lines = content.split("\n")
                if lines[-1].strip() == "```":
                    lines = lines[1:-1]
                else:
                    lines = lines[1:]
                content = "\n".join(lines)

            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            topic_slug = _slugify(topic)[:30]
            ext = "csv" if data_format == "csv" else "json"
            filename = f"{timestamp}-{topic_slug}.{ext}"
            data_dir = project_path / "output" / "data"
            data_dir.mkdir(parents=True, exist_ok=True)
            file_path = data_dir / filename
            file_path.write_text(content, encoding="utf-8")

            yield _sse_event({"type": "content", "text": content, "format": data_format})
            yield _sse_event({"type": "done", "message": f"{data_format.upper()} generated", "file_path": f"output/data/{filename}"})

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        queued_sse_stream(data_stream()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.post("/api/projects/{project_id}/generate-code")
async def generate_code(project_id: str, request: Request):
    """Generate a multi-file code project. Streams SSE."""
    body = await request.json()
    description = (body.get("description") or "").strip()
    if not description:
        return JSONResponse({"error": "No description provided"}, status_code=400)

    project_path = _resolve_project_path(project_id)
    if project_path is None or not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    model = body.get("model", _svc_config.FALLBACK_MODEL)

    async def code_stream():
        try:
            yield _sse_event({"type": "status", "message": f"Generating code: {description[:80]}"})

            prompt = f"""Generate a complete code project for: "{description}"

IMPORTANT: For EACH file, use this EXACT format:
=== FILE: path/to/filename.ext ===
<file content here>
=== END FILE ===

Rules:
- Generate all necessary files for a working project
- Include a README.md with setup instructions
- Use best practices and proper project structure
- Include basic error handling
- Add comments where helpful
- Make sure the code is complete and runnable"""

            yield _sse_event({"type": "status", "message": "Writing code..."})

            content = ""
            async with httpx.AsyncClient(timeout=httpx.Timeout(connect=5.0, read=None, write=120.0, pool=5.0)) as client:
                async with client.stream("POST", f"{_svc_config.OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": True,
                    "options": {"temperature": 0.3, "num_predict": 8192, "num_ctx": 32768},
                }) as resp:
                    async for line in resp.aiter_lines():
                        line = line.strip()
                        if not line:
                            continue
                        try:
                            data = json.loads(line)
                            token = data.get("message", {}).get("content", "")
                            if token:
                                content += token
                                yield _sse_event({"type": "token", "text": token})
                            if data.get("done"):
                                break
                        except json.JSONDecodeError:
                            continue

            if not content:
                yield _sse_event({"type": "error", "error": "Model returned empty content"})
                return

            files = _parse_code_files(content)

            if not files:
                timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
                slug = _slugify(description)[:30]
                filename = f"{timestamp}-{slug}.txt"
                code_dir = project_path / "output" / "code"
                code_dir.mkdir(parents=True, exist_ok=True)
                (code_dir / filename).write_text(content, encoding="utf-8")
                yield _sse_event({"type": "file", "path": f"output/code/{filename}", "size": len(content)})
                yield _sse_event({"type": "done", "message": "Code generated (single file)", "files_count": 1})
                return

            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            slug = _slugify(description)[:30]
            code_dir = project_path / "output" / "code" / f"{timestamp}-{slug}"
            code_dir.mkdir(parents=True, exist_ok=True)

            saved_files = []
            for file_path, file_content in files:
                safe_path = Path(file_path.lstrip("/").replace("..", ""))
                full_path = code_dir / safe_path
                full_path.parent.mkdir(parents=True, exist_ok=True)
                full_path.write_text(file_content, encoding="utf-8")
                rel_path = str(full_path.relative_to(project_path))
                saved_files.append(rel_path)
                yield _sse_event({"type": "file", "path": rel_path, "size": len(file_content)})

            yield _sse_event({"type": "done", "message": f"Generated {len(saved_files)} files", "files_count": len(saved_files)})

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        queued_sse_stream(code_stream()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── Interactive Code Workflow ─────────────────────────────────────────

@router.post("/api/projects/{project_id}/code/plan")
async def code_plan(project_id: str, request: Request):
    """Start an interactive code planning session.

    Phase 1 — Clarification: The AI asks questions to understand the task.
    Phase 2 — Plan: Generates a Markdown plan document.

    Streams SSE events:
      - {type: "question", text: "..."} — clarification question
      - {type: "token", text: "..."} — plan tokens (streaming)
      - {type: "plan", plan_md: "..."} — final plan document
      - {type: "done", session_id: "..."}
    """
    from services.code_session import create_session, update_session, add_conversation_message, save_spec_file
    from services.research import _do_web_search, _fetch_page_content

    body = await request.json()
    description = (body.get("description") or "").strip()
    if not description:
        return JSONResponse({"error": "No description provided"}, status_code=400)

    project_path = _resolve_project_path(project_id)
    if project_path is None or not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    model = body.get("model", _svc_config.FALLBACK_MODEL)
    skip_plan = body.get("skip_plan", False)

    # Create session
    session = create_session(project_id, description)
    add_conversation_message(project_id, "user", f"Build: {description}")

    async def plan_stream():
        try:
            from services.steering import get_steering_context
            steering_ctx = get_steering_context(project_id)

            if skip_plan:
                # Skip straight to a minimal plan
                plan_md = f"# Code Project: {description}\n\nDirect generation (planning skipped)."
                update_session(project_id, status="planned", plan_md=plan_md)
                add_conversation_message(project_id, "assistant", plan_md)
                yield _sse_event({"type": "plan", "plan_md": plan_md})
                yield _sse_event({"type": "done", "session_id": session["id"]})
                return

            # Phase 1: Ask clarifying questions
            yield _sse_event({"type": "status", "message": "Analyzing requirements..."})

            clarify_prompt = f"""You are a senior software architect planning a code project.

The user wants to build: "{description}"

Before writing any code, ask 2-3 specific clarifying questions to ensure you fully understand the requirements. Focus on:
- Technology preferences (framework, language version, etc.)
- Key features or constraints not mentioned
- Any integrations or APIs needed

If the project description is already very specific and clear, you may ask just 1 question.

IMPORTANT: Return ONLY the questions, one per line, numbered. No other text."""

            async with httpx.AsyncClient(timeout=30.0) as client:
                resp = await client.post(f"{_svc_config.OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": clarify_prompt}],
                    "stream": False,
                    "options": {"temperature": 0.5, "num_predict": 512},
                })
                data = resp.json()
                questions_text = data.get("message", {}).get("content", "")

            questions = [q.strip() for q in questions_text.strip().split("\n") if q.strip()]
            questions = [q.lstrip("0123456789.-) ") for q in questions if len(q) > 10][:3]

            if questions:
                update_session(project_id, clarification_questions=questions)
                for q in questions:
                    yield _sse_event({"type": "question", "text": q})
                add_conversation_message(project_id, "assistant", "Questions:\n" + "\n".join(f"- {q}" for q in questions))

            # Phase 2: Optional web research for uncertain technologies
            yield _sse_event({"type": "status", "message": "Researching best practices..."})
            research_context = ""
            try:
                search_results = await _do_web_search(f"best practices {description} code project 2024", max_results=3)
                for r in search_results[:2]:
                    page = await _fetch_page_content(r.get("href", ""), max_chars=2000)
                    if page:
                        research_context += f"\n---\n{page[:1500]}\n"
            except Exception:
                pass

            # Phase 3: Generate the requirements
            yield _sse_event({"type": "status", "message": "Writing requirements..."})

            req_prompt = f"""You are a senior requirements analyst using the EARS (Easy Approach to Requirements Syntax) format.

Project: \"{description}\"

{f"Research context:{research_context}" if research_context else ""}

{f"Steering context:\n{steering_ctx}" if steering_ctx else ""}

Generate a comprehensive requirements.md document with these sections:

# Requirements: {description}

## Overview
Brief description of what will be built and why.

## Requirements

### Requirement 1: [Title]
**User Story:** As a [role], I want [functionality], so that [benefit].

#### Acceptance Criteria
1. WHEN [event] THEN the system SHALL [response]
2. IF [condition] THEN the system SHALL [behavior]

#### Details
- **Priority**: High/Medium/Low
- **Complexity**: High/Medium/Low

(Repeat for each requirement - aim for 4-8 requirements)

## Non-Functional Requirements
### Performance
- WHEN [load condition] THEN the system SHALL [criteria]
### Security
- WHEN [security event] THEN the system SHALL [response]
### Usability
- WHEN [user interaction] THEN the system SHALL [standard]

## Constraints & Assumptions
- Technical constraints
- Business constraints
- Key assumptions

## Success Criteria
- Measurable outcomes that define done

Be specific. Use EARS notation (WHEN/IF/WHILE/WHERE...SHALL) for all acceptance criteria."""

            requirements_md = ""
            async with httpx.AsyncClient(timeout=httpx.Timeout(connect=5.0, read=None, write=120.0, pool=5.0)) as client:
                async with client.stream("POST", f"{_svc_config.OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": req_prompt}],
                    "stream": True,
                    "options": {"temperature": 0.4, "num_predict": 4096, "num_ctx": 32768},
                }) as resp:
                    async for line in resp.aiter_lines():
                        line = line.strip()
                        if not line:
                            continue
                        try:
                            data = json.loads(line)
                            token = data.get("message", {}).get("content", "")
                            if token:
                                requirements_md += token
                                yield _sse_event({"type": "token", "text": token})
                            if data.get("done"):
                                break
                        except json.JSONDecodeError:
                            continue

            if requirements_md:
                update_session(project_id, requirements_md=requirements_md, spec_phase='requirements_review', status='requirements_review', plan_md=requirements_md)
                save_spec_file(project_id, 'requirements.md', requirements_md)
                add_conversation_message(project_id, "assistant", requirements_md)
                yield _sse_event({"type": "plan", "plan_md": requirements_md})

            yield _sse_event({"type": "done", "session_id": session["id"], "spec_phase": "requirements_review"})

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        queued_sse_stream(plan_stream()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.post("/api/projects/{project_id}/code/generate")
async def code_generate_from_plan(project_id: str, request: Request):
    """Generate code from an approved plan. Streams SSE.

    Uses the session's plan_md and conversation context to generate code.
    Updates the session with generated_files and sets status to 'active'.
    """
    from services.code_session import get_session, update_session, add_conversation_message

    body = await request.json()
    session = get_session(project_id)
    if not session:
        return JSONResponse({"error": "No active code session"}, status_code=404)

    project_path = _resolve_project_path(project_id)
    if project_path is None or not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    model = body.get("model", _svc_config.FALLBACK_MODEL)
    # User may provide clarification answers
    answers = body.get("answers", [])
    if answers:
        update_session(project_id, clarification_answers=answers)
        add_conversation_message(project_id, "user", "Answers:\n" + "\n".join(f"- {a}" for a in answers))

    plan_md = session.get("plan_md", "")
    description = session.get("description", "")

    update_session(project_id, status="generating")

    async def gen_stream():
        try:
            yield _sse_event({"type": "status", "message": "Generating code from plan..."})

            answers_context = ""
            if answers:
                answers_context = "\n\nUser's clarification answers:\n" + "\n".join(f"- {a}" for a in answers)

            prompt = f"""Generate a complete code project based on this plan.

Project: "{description}"
{answers_context}

Plan:
{plan_md}

IMPORTANT: For EACH file, use this EXACT format:
=== FILE: path/to/filename.ext ===
<file content here>
=== END FILE ===

Rules:
- Follow the plan's file structure exactly
- Generate ALL files listed in the plan
- Include a README.md with setup and run instructions
- Use best practices, proper error handling, and comments
- Make sure the code is complete and runnable
- Include any config files (package.json, requirements.txt, etc.)"""

            yield _sse_event({"type": "status", "message": "Writing code..."})

            content = ""
            async with httpx.AsyncClient(timeout=httpx.Timeout(connect=5.0, read=None, write=120.0, pool=5.0)) as client:
                async with client.stream("POST", f"{_svc_config.OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": True,
                    "options": {"temperature": 0.3, "num_predict": 8192, "num_ctx": 32768},
                }) as resp:
                    async for line in resp.aiter_lines():
                        line = line.strip()
                        if not line:
                            continue
                        try:
                            data = json.loads(line)
                            token = data.get("message", {}).get("content", "")
                            if token:
                                content += token
                                yield _sse_event({"type": "token", "text": token})
                            if data.get("done"):
                                break
                        except json.JSONDecodeError:
                            continue

            if not content:
                yield _sse_event({"type": "error", "error": "Model returned empty content"})
                return

            files = _parse_code_files(content)
            if not files:
                yield _sse_event({"type": "error", "error": "Could not parse generated files"})
                return

            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            slug = _slugify(description)[:30]
            code_dir = project_path / "output" / "code" / f"{timestamp}-{slug}"
            code_dir.mkdir(parents=True, exist_ok=True)

            saved_files = []
            for file_path, file_content in files:
                safe_path = Path(file_path.lstrip("/").replace("..", ""))
                full_path = code_dir / safe_path
                full_path.parent.mkdir(parents=True, exist_ok=True)
                full_path.write_text(file_content, encoding="utf-8")
                rel_path = str(full_path.relative_to(project_path))
                saved_files.append(rel_path)
                yield _sse_event({"type": "file", "path": rel_path, "size": len(file_content)})

            update_session(project_id, status="active", generated_files=saved_files)
            add_conversation_message(project_id, "assistant", f"Generated {len(saved_files)} files.")

            yield _sse_event({"type": "done", "message": f"Generated {len(saved_files)} files", "files_count": len(saved_files), "session_id": session["id"]})

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        queued_sse_stream(gen_stream()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.post("/api/projects/{project_id}/code/edit")
async def code_edit(project_id: str, request: Request):
    """Edit generated code files based on user instructions.

    Reads the current file contents, sends them with the edit instruction
    to the LLM, writes the changes, and returns a structured change summary.

    Streams SSE:
      - {type: "status", message: "..."}
      - {type: "token", text: "..."} — streamed response
      - {type: "change", file: "...", action: "modified"|"created"}
      - {type: "summary", text: "..."} — human-readable summary
      - {type: "done"}
    """
    from services.code_session import (
        get_session, add_conversation_message, add_edit_record,
        get_session_files_content,
    )
    from services.versions import save_version

    body = await request.json()
    instruction = (body.get("instruction") or "").strip()
    if not instruction:
        return JSONResponse({"error": "No edit instruction provided"}, status_code=400)

    session = get_session(project_id)
    if not session or session.get("status") != "active":
        return JSONResponse({"error": "No active code session. Generate code first."}, status_code=404)

    project_path = _resolve_project_path(project_id)
    if project_path is None or not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    model = body.get("model", _svc_config.FALLBACK_MODEL)
    files_context = get_session_files_content(project_id)

    # Add user instruction to conversation
    add_conversation_message(project_id, "user", instruction)

    # Build conversation context from session history
    conv_messages = session.get("conversation", [])[-10:]

    async def edit_stream():
        try:
            yield _sse_event({"type": "status", "message": f"Analyzing edit request: {instruction[:60]}..."})

            edit_prompt = f"""You are editing an existing code project. The user wants you to make changes.

CURRENT PROJECT FILES:
{files_context}

USER'S EDIT REQUEST: {instruction}

INSTRUCTIONS:
1. Make the requested changes to the relevant files.
2. For EACH file you modify or create, use this EXACT format:
   === FILE: exact/path/to/file.ext ===
   <complete file content with changes applied>
   === END FILE ===

3. After all file blocks, write a CHANGES SUMMARY section:
   === CHANGES SUMMARY ===
   - file1.ext: Description of what changed
   - file2.ext: Description of what changed
   === END SUMMARY ===

4. ONLY output files that were actually changed or newly created.
5. Include the COMPLETE file content for each modified file (not just the diff).
6. Preserve all existing functionality unless the user specifically asks to remove it."""

            messages = [{"role": "user", "content": edit_prompt}]
            # Include relevant conversation context
            if len(conv_messages) > 2:
                context_msg = "Previous conversation context:\n" + "\n".join(
                    f"{m['role']}: {m['content'][:200]}" for m in conv_messages[:-1]
                )
                messages.insert(0, {"role": "system", "content": context_msg})

            yield _sse_event({"type": "status", "message": "Applying changes..."})

            content = ""
            async with httpx.AsyncClient(timeout=httpx.Timeout(connect=5.0, read=None, write=120.0, pool=5.0)) as client:
                async with client.stream("POST", f"{_svc_config.OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": messages,
                    "stream": True,
                    "options": {"temperature": 0.2, "num_predict": 8192, "num_ctx": 32768},
                }) as resp:
                    async for line in resp.aiter_lines():
                        line = line.strip()
                        if not line:
                            continue
                        try:
                            data = json.loads(line)
                            token = data.get("message", {}).get("content", "")
                            if token:
                                content += token
                                yield _sse_event({"type": "token", "text": token})
                            if data.get("done"):
                                break
                        except json.JSONDecodeError:
                            continue

            if not content:
                yield _sse_event({"type": "error", "error": "Model returned empty content"})
                return

            # Parse modified files
            files = _parse_code_files(content)
            changes = []

            for file_rel, file_content in files:
                # Find the matching generated file path
                matched_path = None
                for gf in session.get("generated_files", []):
                    if gf.endswith(file_rel) or file_rel in gf:
                        matched_path = gf
                        break

                if matched_path:
                    full_path = project_path / matched_path
                    action = "modified"
                else:
                    # New file — place in the code directory
                    if session.get("generated_files"):
                        # Use the same code directory as existing files
                        base_dir = str(Path(session["generated_files"][0]).parent)
                        matched_path = f"{base_dir}/{file_rel}"
                    else:
                        matched_path = f"output/code/{file_rel}"
                    full_path = project_path / matched_path
                    action = "created"

                # Save version before overwriting
                if full_path.exists():
                    try:
                        save_version(project_id, matched_path)
                    except Exception:
                        pass

                full_path.parent.mkdir(parents=True, exist_ok=True)
                full_path.write_text(file_content, encoding="utf-8")
                changes.append({"file": matched_path, "action": action})
                yield _sse_event({"type": "change", "file": matched_path, "action": action})

                # Update generated_files if new file
                if action == "created" and matched_path not in session.get("generated_files", []):
                    gf = session.get("generated_files", [])
                    gf.append(matched_path)
                    from services.code_session import update_session
                    update_session(project_id, generated_files=gf)

            # Execute matching hooks in background
            try:
                from services.hooks import evaluate_hooks, execute_hook
                import asyncio as _asyncio
                for changed in changes:
                    evt_type = 'file_saved' if changed['action'] == 'modified' else 'file_created'
                    matching = evaluate_hooks(project_id, evt_type, changed.get('file', ''))
                    for h in matching:
                        _asyncio.create_task(execute_hook(project_id, h['id'], f"File {changed['action']}: {changed.get('file', '')}", model))
            except Exception:
                pass  # Hooks are best-effort

            # Parse change summary
            summary_match = re.search(
                r'===\s*CHANGES?\s*SUMMARY\s*===\s*\n(.*?)\n===\s*END\s*SUMMARY\s*===',
                content, re.DOTALL
            )
            summary = summary_match.group(1).strip() if summary_match else (
                "\n".join(f"- {c['file']}: {c['action']}" for c in changes)
            )

            add_conversation_message(project_id, "assistant", f"Changes applied:\n{summary}")
            add_edit_record(project_id, instruction, changes, summary)

            yield _sse_event({"type": "summary", "text": summary, "changes": changes})
            yield _sse_event({"type": "done", "message": f"Applied {len(changes)} file changes"})

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        queued_sse_stream(edit_stream()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.get("/api/projects/{project_id}/code/session")
async def get_code_session(project_id: str):
    """Get the current code session for a project."""
    from services.code_session import get_session
    session = get_session(project_id)
    if not session:
        return JSONResponse({"error": "No active code session"}, status_code=404)
    return session


@router.delete("/api/projects/{project_id}/code/session")
async def close_code_session(project_id: str):
    """Close and delete the current code session."""
    from services.code_session import close_session, delete_session
    close_session(project_id)
    delete_session(project_id)
    return {"ok": True}


@router.post("/api/projects/{project_id}/workflow")
async def run_workflow(project_id: str, request: Request):
    """Execute a multi-step workflow. Streams SSE progress."""
    body = await request.json()
    user_request = (body.get("request") or "").strip()
    if not user_request:
        return JSONResponse({"error": "No request provided"}, status_code=400)

    project_path = _resolve_project_path(project_id)
    if project_path is None or not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    model = body.get("model", _svc_config.FALLBACK_MODEL)

    async def workflow_stream():
        try:
            yield _sse_event({"type": "status", "message": "Planning workflow..."})

            plan_prompt = f"""Analyze this request and break it into ordered steps. For each step, specify the type.

Request: "{user_request}"

Available step types:
- research: Search the web and gather information on a topic
- document: Write a document/report about a topic
- code: Generate code for a programming task
- data: Generate structured data (CSV/JSON tables)

Return ONLY a JSON array of steps, no explanation. Example:
[{{"type": "research", "description": "Research topic X"}}, {{"type": "document", "description": "Write a report on X"}}]"""

            async with httpx.AsyncClient(timeout=httpx.Timeout(connect=5.0, read=None, write=120.0, pool=5.0)) as client:
                async with client.stream("POST", f"{_svc_config.OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": plan_prompt}],
                    "stream": True,
                    "options": {"temperature": 0.2, "num_predict": 1024},
                }) as resp:
                    plan_text = ""
                    async for line in resp.aiter_lines():
                        line = line.strip()
                        if not line:
                            continue
                        try:
                            data = json.loads(line)
                            token = data.get("message", {}).get("content", "")
                            if token:
                                plan_text += token
                            if data.get("done"):
                                break
                        except json.JSONDecodeError:
                            continue

            steps = _parse_workflow_plan(plan_text)
            if not steps:
                steps = [{"type": "document", "description": user_request}]

            yield _sse_event({"type": "plan", "steps": steps})

            for i, step in enumerate(steps):
                step_type = step.get("type", "document")
                step_desc = step.get("description", user_request)

                yield _sse_event({
                    "type": "step_start",
                    "step": i + 1,
                    "total": len(steps),
                    "step_type": step_type,
                    "description": step_desc,
                })

                if step_type == "research":
                    queries = await _generate_search_queries(step_desc, 3, model)
                    all_sources = []
                    all_content = []
                    for query in queries:
                        results = await _do_web_search(query, max_results=3)
                        for r in results[:2]:
                            all_sources.append(r)
                            page = await _fetch_page_content(r.get("href", ""))
                            if page:
                                all_content.append(page[:2000])

                    if all_content:
                        findings = await _extract_findings(step_desc, all_content, model)
                        knowledge_file = project_path / "knowledge.json"
                        knowledge = json.loads(knowledge_file.read_text(encoding="utf-8"))
                        knowledge.setdefault("findings", []).append({
                            "topic": step_desc,
                            "summary": findings[:2000],
                            "sources": [s.get("href", "") for s in all_sources[:10]],
                            "timestamp": datetime.now(timezone.utc).isoformat(),
                        })
                        knowledge_file.write_text(json.dumps(knowledge, indent=2, ensure_ascii=False), encoding="utf-8")

                    yield _sse_event({"type": "step_done", "step": i + 1, "message": f"Research complete: {len(all_sources)} sources found"})

                elif step_type == "document":
                    context = _get_project_knowledge_context(project_id, max_chars=6000)
                    doc_prompt = f"Write a comprehensive document about: {step_desc}\n\n{context if context else ''}"

                    doc_content = ""
                    async with httpx.AsyncClient(timeout=httpx.Timeout(connect=5.0, read=None, write=120.0, pool=5.0)) as client:
                        async with client.stream("POST", f"{_svc_config.OLLAMA}/api/chat", json={
                            "model": model,
                            "messages": [{"role": "user", "content": doc_prompt}],
                            "stream": True,
                            "options": {"temperature": 0.5, "num_predict": 8192, "num_ctx": 32768},
                        }) as resp:
                            async for line in resp.aiter_lines():
                                line = line.strip()
                                if not line:
                                    continue
                                try:
                                    data = json.loads(line)
                                    token = data.get("message", {}).get("content", "")
                                    if token:
                                        doc_content += token
                                        yield _sse_event({"type": "step_token", "step": i + 1, "text": token})
                                    if data.get("done"):
                                        break
                                except json.JSONDecodeError:
                                    continue

                    if doc_content:
                        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
                        slug = _slugify(step_desc)[:30]
                        output_dir = project_path / "output"
                        output_dir.mkdir(exist_ok=True)
                        (output_dir / f"{timestamp}-{slug}.md").write_text(doc_content, encoding="utf-8")

                    yield _sse_event({"type": "step_done", "step": i + 1, "message": "Document generated"})

                elif step_type == "code":
                    code_prompt = f"""Generate a complete code project for: "{step_desc}"

For EACH file use: === FILE: path === ... === END FILE ==="""

                    code_content = ""
                    async with httpx.AsyncClient(timeout=httpx.Timeout(connect=5.0, read=None, write=120.0, pool=5.0)) as client:
                        async with client.stream("POST", f"{_svc_config.OLLAMA}/api/chat", json={
                            "model": model,
                            "messages": [{"role": "user", "content": code_prompt}],
                            "stream": True,
                            "options": {"temperature": 0.3, "num_predict": 8192, "num_ctx": 32768},
                        }) as resp:
                            async for line in resp.aiter_lines():
                                line = line.strip()
                                if not line:
                                    continue
                                try:
                                    data = json.loads(line)
                                    token = data.get("message", {}).get("content", "")
                                    if token:
                                        code_content += token
                                        yield _sse_event({"type": "step_token", "step": i + 1, "text": token})
                                    if data.get("done"):
                                        break
                                except json.JSONDecodeError:
                                    continue

                    if code_content:
                        files = _parse_code_files(code_content)
                        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
                        slug = _slugify(step_desc)[:30]
                        code_dir = project_path / "output" / "code" / f"{timestamp}-{slug}"
                        code_dir.mkdir(parents=True, exist_ok=True)
                        if files:
                            for fp, fc in files:
                                safe = Path(fp.lstrip("/").replace("..", ""))
                                full = code_dir / safe
                                full.parent.mkdir(parents=True, exist_ok=True)
                                full.write_text(fc, encoding="utf-8")
                        else:
                            (code_dir / "output.txt").write_text(code_content, encoding="utf-8")

                    yield _sse_event({"type": "step_done", "step": i + 1, "message": "Code generated"})

                elif step_type == "data":
                    data_prompt = f"""Generate a CSV dataset about: "{step_desc}"
Return ONLY CSV content with headers. No explanations."""

                    csv_content = ""
                    async with httpx.AsyncClient(timeout=httpx.Timeout(connect=5.0, read=None, write=120.0, pool=5.0)) as client:
                        async with client.stream("POST", f"{_svc_config.OLLAMA}/api/chat", json={
                            "model": model,
                            "messages": [{"role": "user", "content": data_prompt}],
                            "stream": True,
                            "options": {"temperature": 0.3, "num_predict": 4096, "num_ctx": 32768},
                        }) as resp:
                            async for line in resp.aiter_lines():
                                line = line.strip()
                                if not line:
                                    continue
                                try:
                                    data = json.loads(line)
                                    token = data.get("message", {}).get("content", "")
                                    if token:
                                        csv_content += token
                                        yield _sse_event({"type": "step_token", "step": i + 1, "text": token})
                                    if data.get("done"):
                                        break
                                except json.JSONDecodeError:
                                    continue

                    if csv_content:
                        csv_content = csv_content.strip()
                        if csv_content.startswith("```"):
                            lines = csv_content.split("\n")
                            if lines[-1].strip() == "```":
                                lines = lines[1:-1]
                            else:
                                lines = lines[1:]
                            csv_content = "\n".join(lines)

                        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
                        slug = _slugify(step_desc)[:30]
                        data_dir = project_path / "output" / "data"
                        data_dir.mkdir(parents=True, exist_ok=True)
                        (data_dir / f"{timestamp}-{slug}.csv").write_text(csv_content, encoding="utf-8")

                    yield _sse_event({"type": "step_done", "step": i + 1, "message": "Data generated"})

                else:
                    yield _sse_event({"type": "step_done", "step": i + 1, "message": f"Skipped unknown step type: {step_type}"})

            yield _sse_event({"type": "done", "message": f"Workflow complete! Executed {len(steps)} steps."})

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        queued_sse_stream(workflow_stream()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )
