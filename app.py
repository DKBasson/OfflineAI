import asyncio
import atexit
import json
import os
import platform
import re
import shlex
import shutil
import subprocess
import tempfile
import threading
import time
import unicodedata
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from fastapi import FastAPI, File, Request, UploadFile
from fastapi.responses import HTMLResponse, StreamingResponse, JSONResponse, Response, FileResponse
from fastapi.staticfiles import StaticFiles
import httpx
import logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("offlineai")

try:
    from ddgs import DDGS as _DDGS
    _SEARCH_AVAILABLE = True
except ImportError:
    try:
        from duckduckgo_search import DDGS as _DDGS
        _SEARCH_AVAILABLE = True
    except ImportError:
        _SEARCH_AVAILABLE = False

try:
    from faster_whisper import WhisperModel as _WhisperModel
    _WHISPER_AVAILABLE = True
except ImportError:
    _WHISPER_AVAILABLE = False

try:
    import docx as _docx_module
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    from odf.opendocument import load as _odf_load
    from odf.teletype import extractText as _odf_extract_text
    from odf import text as _odf_text
    _ODF_AVAILABLE = True
except ImportError:
    _ODF_AVAILABLE = False

try:
    import pypdf as _pypdf
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    _BS4_AVAILABLE = True
except ImportError:
    _BS4_AVAILABLE = False

try:
    import markdown as _markdown_lib
    _MARKDOWN_AVAILABLE = True
except ImportError:
    _MARKDOWN_AVAILABLE = False

try:
    import weasyprint as _weasyprint
    _WEASYPRINT_AVAILABLE = True
except ImportError:
    _WEASYPRINT_AVAILABLE = False

import image_gen as _image_gen

_WHISPER_MODEL_SIZE = os.environ.get("WHISPER_MODEL", "tiny")
_whisper_model: object = None
_whisper_lock = asyncio.Lock()


async def _get_whisper() -> object:
    global _whisper_model
    async with _whisper_lock:
        if _whisper_model is None:
            _whisper_model = await asyncio.to_thread(
                _WhisperModel, _WHISPER_MODEL_SIZE, device="cpu", compute_type="int8"
            )
    return _whisper_model

app = FastAPI(title="OfflineAI")
_SERVER_START_TIME = time.time()

OLLAMA     = os.environ.get("OLLAMA_URL", "http://localhost:11434").rstrip("/")
HOST       = os.environ.get("OFFLINEAI_HOST", "127.0.0.1")
PORT       = int(os.environ.get("OFFLINEAI_PORT", "8080"))
AUTH_TOKEN = os.environ.get("OFFLINEAI_TOKEN", "").strip()
LAN_MODE   = HOST in {"0.0.0.0", "::"}
AUTH_REQUIRED = LAN_MODE and bool(AUTH_TOKEN)
LOOPBACK_HOSTS = {"127.0.0.1", "::1", "localhost"}
STATIC_DIR = Path(__file__).parent / "static"
FRONTEND_DIR = Path(__file__).parent / "frontend"
MAX_BODY   = 50 * 1024 * 1024
FALLBACK_MODEL = "gemma4:e4b"
IMAGE_GEN_MAX_WIDTH = int(os.environ.get("OFFLINEAI_IMAGE_MAX_WIDTH", "1024"))
IMAGE_GEN_MAX_HEIGHT = int(os.environ.get("OFFLINEAI_IMAGE_MAX_HEIGHT", "1024"))
IMAGE_GEN_MAX_STEPS = int(os.environ.get("OFFLINEAI_IMAGE_MAX_STEPS", "16"))
IMAGE_GEN_DEFAULT_WIDTH = int(os.environ.get("OFFLINEAI_IMAGE_DEFAULT_WIDTH", "640"))
IMAGE_GEN_DEFAULT_HEIGHT = int(os.environ.get("OFFLINEAI_IMAGE_DEFAULT_HEIGHT", "640"))
IMAGE_GEN_DEFAULT_STEPS = int(os.environ.get("OFFLINEAI_IMAGE_DEFAULT_STEPS", "6"))

_TOKEN_STATS_FILE = Path(__file__).parent / "token_stats.json"
_MAX_TOKEN_STATS_ENTRIES = int(os.environ.get("OFFLINEAI_MAX_TOKEN_ENTRIES", "500"))

def _load_token_stats() -> dict:
    try:
        if _TOKEN_STATS_FILE.exists():
            raw = json.loads(_TOKEN_STATS_FILE.read_text(encoding="utf-8"))
            return {k: v for k, v in raw.items()
                    if isinstance(v, list) and len(v) == 2
                    and all(isinstance(x, (int, float)) for x in v)}
    except Exception:
        pass
    return {}

def _save_token_stats() -> None:
    global _token_stats
    if len(_token_stats) > _MAX_TOKEN_STATS_ENTRIES:
        _token_stats = dict(
            sorted(
                _token_stats.items(),
                key=lambda x: x[1][0] + x[1][1],
                reverse=True,
            )[:_MAX_TOKEN_STATS_ENTRIES]
        )
    try:
        _TOKEN_STATS_FILE.write_text(json.dumps(_token_stats), encoding="utf-8")
    except Exception:
        pass

_token_stats: dict[str, list[int]] = _load_token_stats()

_BASE = Path(__file__).parent
_STYLES_CSS = (_BASE / "styles.css").read_text(encoding="utf-8")
_REACT_DIST = _BASE / "react-dist"
_REACT_MODE = _REACT_DIST.is_dir() and (_REACT_DIST / "index.html").is_file()

if STATIC_DIR.is_dir():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

if _REACT_MODE:
    _react_assets = _REACT_DIST / "assets"
    if _react_assets.is_dir():
        app.mount("/assets", StaticFiles(directory=str(_react_assets)), name="react-assets")
elif FRONTEND_DIR.is_dir():
    app.mount("/frontend", StaticFiles(directory=str(FRONTEND_DIR)), name="frontend")


@app.middleware("http")
async def require_lan_token(request: Request, call_next):
    client_host = request.client.host if request.client else ""
    is_loopback = client_host in LOOPBACK_HOSTS
    if AUTH_REQUIRED and not is_loopback and request.url.path.startswith("/api/"):
        supplied = request.headers.get("x-offlineai-token") or request.query_params.get("token")
        if supplied != AUTH_TOKEN:
            log.warning("Auth rejected for %s on %s", client_host, request.url.path)
            return JSONResponse({"error": "Authentication required"}, status_code=401)
    return await call_next(request)


@app.middleware("http")
async def limit_body_size(request: Request, call_next):
    content_length = request.headers.get("content-length")
    try:
        body_size = int(content_length) if content_length else 0
    except ValueError:
        return JSONResponse({"error": "Invalid Content-Length header"}, status_code=400)
    if body_size > MAX_BODY:
        log.warning("Request too large: %s bytes", body_size)
        return JSONResponse({"error": "Request body too large (max 50 MB)"}, status_code=413)
    return await call_next(request)


def _is_loopback_request(request: Request) -> bool:
    client_host = request.client.host if request.client else ""
    return client_host in LOOPBACK_HOSTS or client_host == "::ffff:127.0.0.1"


def _runtime_control_allowed(request: Request) -> bool:
    if _is_loopback_request(request):
        return True
    supplied = request.headers.get("x-offlineai-token") or request.query_params.get("token")
    return bool(AUTH_TOKEN and supplied == AUTH_TOKEN)


def _start_ollama_serve() -> tuple[bool, str]:
    ollama = shutil.which("ollama")
    if not ollama:
        return False, "The ollama command was not found on PATH."

    kwargs = {
        "stdin": subprocess.DEVNULL,
        "stdout": subprocess.DEVNULL,
        "stderr": subprocess.DEVNULL,
    }
    if platform.system() == "Windows":
        kwargs["creationflags"] = subprocess.CREATE_NEW_PROCESS_GROUP | subprocess.DETACHED_PROCESS
    else:
        kwargs["start_new_session"] = True
    subprocess.Popen([ollama, "serve"], **kwargs)
    return True, "ollama serve started"


def _restart_ollama_process() -> dict:
    custom_cmd = os.environ.get("OLLAMA_RESTART_CMD", "").strip()
    if custom_cmd:
        try:
            cmd = shlex.split(custom_cmd)
            if not cmd:
                return {"ok": False, "error": "OLLAMA_RESTART_CMD is empty."}
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=30, check=False)
            if proc.returncode != 0:
                detail = (proc.stderr or proc.stdout or "").strip()
                return {
                    "ok": False,
                    "error": f"Restart command exited with {proc.returncode}.",
                    "detail": detail,
                    "method": "custom",
                }
            return {"ok": True, "method": "custom"}
        except Exception as exc:
            return {"ok": False, "error": str(exc), "method": "custom"}

    try:
        system = platform.system()
        if system == "Windows":
            taskkill = shutil.which("taskkill")
            if taskkill:
                subprocess.run(
                    [taskkill, "/f", "/im", "ollama.exe"],
                    capture_output=True,
                    text=True,
                    timeout=8,
                    check=False,
                )
        else:
            pkill = shutil.which("pkill")
            killall = shutil.which("killall")
            if pkill:
                subprocess.run([pkill, "-x", "ollama"], capture_output=True, text=True, timeout=8, check=False)
            elif killall:
                subprocess.run([killall, "ollama"], capture_output=True, text=True, timeout=8, check=False)

        time.sleep(0.8)
        started, message = _start_ollama_serve()
        if not started:
            return {"ok": False, "error": message, "method": "ollama serve"}
        return {"ok": True, "message": message, "method": "ollama serve"}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "method": "ollama serve"}


async def _wait_for_ollama_ready(timeout: float = 12.0) -> tuple[bool, str]:
    deadline = time.monotonic() + timeout
    last_error = ""
    async with httpx.AsyncClient(timeout=1.0) as client:
        while time.monotonic() < deadline:
            try:
                resp = await client.get(f"{OLLAMA}/")
                if resp.status_code < 500:
                    return True, ""
                last_error = resp.text
            except Exception as exc:
                last_error = str(exc)
            await asyncio.sleep(0.5)
    return False, last_error


@app.get("/", response_class=HTMLResponse)
async def root():
    if _REACT_MODE:
        return HTMLResponse((_REACT_DIST / "index.html").read_text(encoding="utf-8"))
    return HTMLResponse((_BASE / "index.html").read_text(encoding="utf-8"))


@app.get("/styles.css")
async def styles():
    return Response(_STYLES_CSS, media_type="text/css")


async def _ollama_json_request(method: str, path: str, *, body: dict | None = None, timeout: float = 5.0) -> dict:
    async with httpx.AsyncClient(timeout=timeout) as client:
        response = await client.request(method, f"{OLLAMA}{path}", json=body)
        response.raise_for_status()
        return response.json()


@app.get("/api/models")
async def get_models():
    try:
        return await _ollama_json_request("GET", "/api/tags", timeout=3.0)
    except Exception as exc:
        return {"models": [{"name": FALLBACK_MODEL}], "offline": True, "error": str(exc)}


@app.get("/api/status")
async def status():
    try:
        data = await _ollama_json_request("GET", "/api/tags", timeout=2.0)
        return {
            "ollama": True,
            "models_count": len(data.get("models", [])),
            "lan": HOST in {"0.0.0.0", "::"},
            "auth_required": AUTH_REQUIRED,
            "host": HOST,
            "port": PORT,
        }
    except Exception as exc:
        return JSONResponse(
            {
                "ollama": False,
                "error": str(exc),
                "lan": HOST in {"0.0.0.0", "::"},
                "auth_required": AUTH_REQUIRED,
                "host": HOST,
                "port": PORT,
            },
            status_code=503,
        )


def _get_system_memory() -> tuple[float, float]:
    """Return (total_gb, available_gb) using OS-native methods. No psutil needed."""
    total_gb = 0.0
    avail_gb = 0.0
    system = platform.system()
    try:
        if system == "Darwin":
            raw = subprocess.check_output(["sysctl", "-n", "hw.memsize"], text=True).strip()
            total_bytes = int(raw)
            total_gb = round(total_bytes / (1024 ** 3), 1)
            vm = subprocess.check_output(["vm_stat"], text=True)
            page_size = 4096
            ps_match = re.search(r"page size of (\d+) bytes", vm)
            if ps_match:
                page_size = int(ps_match.group(1))
            free = int(re.search(r"Pages free:\s+(\d+)", vm).group(1)) if re.search(r"Pages free:\s+(\d+)", vm) else 0
            inactive = int(re.search(r"Pages inactive:\s+(\d+)", vm).group(1)) if re.search(r"Pages inactive:\s+(\d+)", vm) else 0
            avail_gb = round((free + inactive) * page_size / (1024 ** 3), 1)
        elif system == "Linux":
            with open("/proc/meminfo") as f:
                meminfo = f.read()
            mt = re.search(r"MemTotal:\s+(\d+)\s+kB", meminfo)
            ma = re.search(r"MemAvailable:\s+(\d+)\s+kB", meminfo)
            if mt:
                total_gb = round(int(mt.group(1)) / (1024 ** 2), 1)
            if ma:
                avail_gb = round(int(ma.group(1)) / (1024 ** 2), 1)
        else:
            total_gb = round(os.sysconf("SC_PAGE_SIZE") * os.sysconf("SC_PHYS_PAGES") / (1024 ** 3), 1) if hasattr(os, "sysconf") else 0
    except Exception:
        pass
    return total_gb, avail_gb


def _dir_size_mb(path: Path) -> float:
    """Return directory size in MB."""
    total = 0
    try:
        for f in path.rglob("*"):
            if f.is_file():
                total += f.stat().st_size
    except Exception:
        pass
    return round(total / (1024 * 1024), 1)


@app.get("/api/health")
async def health_check():
    """Comprehensive health/diagnostics endpoint."""
    # --- Ollama ---
    ollama_info: dict = {"status": "offline", "version": None, "models_count": 0}
    try:
        async with httpx.AsyncClient(timeout=2.0) as client:
            ver_resp = await client.get(f"{OLLAMA}/api/version")
            if ver_resp.status_code == 200:
                ver_data = ver_resp.json()
                ollama_info["version"] = ver_data.get("version")
            tags_resp = await client.get(f"{OLLAMA}/api/tags")
            if tags_resp.status_code == 200:
                ollama_info["status"] = "online"
                ollama_info["models_count"] = len(tags_resp.json().get("models", []))
    except Exception:
        pass

    # --- System ---
    ram_total, ram_avail = await asyncio.to_thread(_get_system_memory)
    try:
        disk = shutil.disk_usage(Path.home())
        disk_free_gb = round(disk.free / (1024 ** 3), 1)
    except Exception:
        disk_free_gb = 0

    system_info = {
        "platform": platform.system().lower(),
        "python": platform.python_version(),
        "ram_total_gb": ram_total,
        "ram_available_gb": ram_avail,
        "disk_free_gb": disk_free_gb,
    }

    # --- Services ---
    services: dict = {}

    # Whisper
    if _WHISPER_AVAILABLE:
        services["whisper"] = {"available": True, "model": _WHISPER_MODEL_SIZE}
    else:
        services["whisper"] = {"available": False, "reason": "faster-whisper not installed"}

    # WeasyPrint
    if _WEASYPRINT_AVAILABLE:
        services["weasyprint"] = {"available": True}
    else:
        services["weasyprint"] = {"available": False, "reason": "weasyprint not installed"}

    # Diffusers / image generation — check if torch is importable
    try:
        import torch  # noqa: F401
        services["diffusers"] = {"available": True}
    except ImportError:
        services["diffusers"] = {"available": False, "reason": "torch not installed"}

    # Document parsing
    if _DOCX_AVAILABLE:
        services["docx"] = {"available": True}
    else:
        services["docx"] = {"available": False, "reason": "python-docx not installed"}

    # --- Projects ---
    project_count = 0
    if PROJECTS_DIR.is_dir():
        project_count = sum(1 for d in PROJECTS_DIR.iterdir() if d.is_dir() and (d / "knowledge.json").is_file())
    projects_disk = await asyncio.to_thread(_dir_size_mb, PROJECTS_DIR) if PROJECTS_DIR.is_dir() else 0

    # --- Tools ---
    registry = _load_tool_registry()
    tools_enabled = sum(1 for t in registry if t.get("enabled", True))
    tools_disabled = len(registry) - tools_enabled

    # --- Memory ---
    memories = _load_memories()

    # --- Uptime ---
    uptime = round(time.time() - _SERVER_START_TIME)

    return {
        "ollama": ollama_info,
        "system": system_info,
        "services": services,
        "projects": {"count": project_count, "disk_usage_mb": projects_disk},
        "tools": {"count": len(registry), "enabled": tools_enabled, "disabled": tools_disabled},
        "memory": {"count": len(memories)},
        "uptime_seconds": uptime,
    }


@app.post("/api/ollama/restart")
async def restart_ollama(request: Request):
    if not _runtime_control_allowed(request):
        return JSONResponse({"error": "Ollama restart requires localhost access or a valid LAN token."}, status_code=403)

    result = await asyncio.to_thread(_restart_ollama_process)
    if not result.get("ok"):
        return JSONResponse(result, status_code=500)

    ready, error = await _wait_for_ollama_ready()
    if not ready:
        return JSONResponse(
            {
                "ok": False,
                "error": "Restart command ran, but Ollama did not respond within 12 seconds.",
                "detail": error,
                "method": result.get("method"),
            },
            status_code=503,
        )

    return {
        "ok": True,
        "message": "Ollama restarted",
        "method": result.get("method"),
    }


@app.post("/api/show")
async def show_model(request: Request):
    body = await request.json()
    try:
        return await _ollama_json_request("POST", "/api/show", body=body, timeout=5.0)
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=502)


@app.post("/api/transcribe")
async def transcribe_audio(file: UploadFile = File(...)):
    """Stream transcription progress for an audio file using faster-whisper (SSE)."""
    if not _WHISPER_AVAILABLE:
        return JSONResponse(
            {"error": "Audio transcription requires faster-whisper. Install with: pip install faster-whisper"},
            status_code=501,
        )
    suffix = Path(file.filename).suffix if file.filename else ".wav"
    content = await file.read()
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    model = await _get_whisper()

    async def event_stream():
        q: asyncio.Queue = asyncio.Queue()
        loop = asyncio.get_running_loop()

        def _run():
            try:
                segments_iter, info = model.transcribe(tmp_path, beam_size=5)
                duration = max(info.duration or 0, 0.001)
                texts: list[str] = []
                for seg in segments_iter:
                    texts.append(seg.text)
                    pct = min(99, int(seg.end / duration * 100))
                    loop.call_soon_threadsafe(
                        q.put_nowait, {"type": "progress", "percent": pct}
                    )
                loop.call_soon_threadsafe(
                    q.put_nowait,
                    {"type": "done", "transcript": " ".join(texts).strip()},
                )
            except Exception as exc:
                loop.call_soon_threadsafe(
                    q.put_nowait, {"type": "error", "error": str(exc)}
                )
            finally:
                Path(tmp_path).unlink(missing_ok=True)

        threading.Thread(target=_run, daemon=True).start()

        while True:
            item = await q.get()
            yield f"data: {json.dumps(item)}\n\n"
            if item["type"] in ("done", "error"):
                break

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/api/extract")
async def extract_document(file: UploadFile = File(...)):
    """Extract plain text from a Word (.docx) or ODF (.odt/.ods/.odp) document."""
    suffix = Path(file.filename).suffix.lower() if file.filename else ""
    content = await file.read()

    if suffix == ".docx":
        if not _DOCX_AVAILABLE:
            return JSONResponse(
                {"error": "python-docx is required. Install with: pip install python-docx"},
                status_code=501,
            )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            def _read_docx():
                doc = _docx_module.Document(tmp_path)
                return "\n".join(p.text for p in doc.paragraphs)
            text = await asyncio.to_thread(_read_docx)
            return {"text": text}
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    elif suffix in {".odt", ".ods", ".odp"}:
        if not _ODF_AVAILABLE:
            return JSONResponse(
                {"error": "odfpy is required. Install with: pip install odfpy"},
                status_code=501,
            )
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            def _read_odf():
                odoc = _odf_load(tmp_path)
                paragraphs = odoc.getElementsByType(_odf_text.P)
                return "\n".join(_odf_extract_text(p) for p in paragraphs)
            text = await asyncio.to_thread(_read_odf)
            return {"text": text}
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    elif suffix == ".pdf":
        if not _PDF_AVAILABLE:
            return JSONResponse(
                {"error": "pypdf is required. Install with: pip install pypdf"},
                status_code=501,
            )
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            def _read_pdf():
                reader = _pypdf.PdfReader(tmp_path)
                return "\n".join(
                    page.extract_text() or "" for page in reader.pages
                ).strip()
            text = await asyncio.to_thread(_read_pdf)
            return {"text": text}
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    else:
        return JSONResponse(
            {"error": f"Unsupported document format: '{suffix}'. Supported: .docx, .odt, .ods, .odp, .pdf"},
            status_code=415,
        )


@app.post("/api/search")
async def web_search(request: Request):
    """Perform a web search using DuckDuckGo and return results."""
    if not _SEARCH_AVAILABLE:
        return JSONResponse(
            {"error": "Web search requires duckduckgo-search. Install with: pip install duckduckgo-search"},
            status_code=501,
        )
    body = await request.json()
    query = (body.get("query") or "").strip()
    if not query:
        return JSONResponse({"error": "No search query provided"}, status_code=400)

    max_results = min(int(body.get("max_results", 5)), 10)

    def _do_search():
        try:
            ddgs = _DDGS()
            results = list(ddgs.text(query, max_results=max_results))
            return results
        except Exception as exc:
            return {"error": str(exc)}

    results = await asyncio.to_thread(_do_search)
    if isinstance(results, dict) and "error" in results:
        return JSONResponse({"error": results["error"]}, status_code=502)

    return {"query": query, "results": results}


def _ndjson_error(message: str) -> bytes:
    return (json.dumps({"error": message}) + "\n").encode()


def _clamp_int(value, minimum: int, maximum: int, fallback: int) -> int:
    try:
        n = int(value)
    except (TypeError, ValueError):
        n = fallback
    return max(minimum, min(maximum, n))


def _apply_image_generation_caps(body: dict) -> dict:
    capped = dict(body or {})
    capped["width"] = _clamp_int(capped.get("width"), 256, IMAGE_GEN_MAX_WIDTH, IMAGE_GEN_DEFAULT_WIDTH)
    capped["height"] = _clamp_int(capped.get("height"), 256, IMAGE_GEN_MAX_HEIGHT, IMAGE_GEN_DEFAULT_HEIGHT)
    capped["steps"] = _clamp_int(capped.get("steps"), 2, IMAGE_GEN_MAX_STEPS, IMAGE_GEN_DEFAULT_STEPS)
    return capped


def _sse_event(data: dict) -> str:
    return f"data: {json.dumps(data)}\n\n"


async def _generate_search_queries(topic: str, num_queries: int, model: str) -> list[str]:
    """Use LLM to generate diverse search queries for a research topic."""
    prompt = f"""Generate exactly {num_queries} diverse web search queries to research the topic: "{topic}"

Rules:
- Each query should explore a different angle or aspect of the topic
- Queries should be specific enough to get relevant results
- Include a mix of overview queries and specific detail queries
- Return ONLY the queries, one per line, no numbering, no explanations"""

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(f"{OLLAMA}/api/chat", json={
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "options": {"temperature": 0.7, "num_predict": 512},
            })
            data = resp.json()
            content = data.get("message", {}).get("content", "")
            queries = [q.strip().strip('"').strip("'") for q in content.strip().split("\n") if q.strip()]
            queries = [re.sub(r'^[\d]+[.)\s]+|^[-*]\s+', '', q).strip() for q in queries]
            return queries[:num_queries] if queries else [topic]
    except Exception:
        return [topic]


async def _do_web_search(query: str, max_results: int = 5) -> list[dict]:
    """Perform a web search using DuckDuckGo."""
    if not _SEARCH_AVAILABLE:
        return []
    try:
        def _search():
            ddgs = _DDGS()
            return list(ddgs.text(query, max_results=max_results))
        return await asyncio.to_thread(_search)
    except Exception:
        return []


async def _fetch_page_content(url: str, max_chars: int = 4000) -> str:
    """Fetch and extract text content from a URL."""
    if not _BS4_AVAILABLE:
        return ""
    try:
        from bs4 import BeautifulSoup
        async with httpx.AsyncClient(timeout=10.0, follow_redirects=True) as client:
            resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0"})
            if resp.status_code != 200:
                return ""
        soup = BeautifulSoup(resp.text, "lxml")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        return text[:max_chars]
    except Exception:
        return ""


async def _extract_findings(topic: str, page_contents: list[str], model: str) -> str:
    """Use LLM to extract key findings from collected source material."""
    if not page_contents:
        return "No source content available to analyze."

    combined = "\n\n---\n\n".join(page_contents[:10])
    combined = combined[:24000]

    prompt = f"""Analyze the following source material about "{topic}" and extract the key findings.

Source material:
{combined}

Provide a structured list of key findings, facts, and insights. Be specific and factual. Include relevant data points, dates, and names where available.
For each finding, note which source it came from by including the source name in parentheses at the end, e.g. "Finding text (Source: Article Title)"."""

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(f"{OLLAMA}/api/chat", json={
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "options": {"temperature": 0.3, "num_predict": 4096, "num_ctx": 32768},
            })
            data = resp.json()
            return data.get("message", {}).get("content", "Unable to extract findings.")
    except Exception as exc:
        return f"Error extracting findings: {exc}"


async def _synthesize_summary(topic: str, findings: str, sources: list[dict], model: str) -> str:
    """Use LLM to write a comprehensive research summary."""
    sources_list = "\n".join(f"[{i}] {s.get('title', 'Unknown')}: {s.get('url', '')}" for i, s in enumerate(sources[:15], 1))

    prompt = f"""Write a comprehensive research summary about "{topic}" based on the following findings and sources.

Key Findings:
{findings}

Numbered Sources:
{sources_list}

Write a well-structured Markdown document with:
1. A title (# heading)
2. An executive summary paragraph
3. Key findings organized by theme (## subheadings)
4. A "## References" section at the end listing all cited sources as a numbered list

CITATION RULES (you MUST follow these):
- Use inline citation markers like [1], [2], etc. to reference the numbered sources above.
- Place the citation marker immediately after the claim or fact it supports.
- Every factual claim should have at least one citation.
- In the References section, list each cited source as: [N] Title — URL
- Only cite sources from the numbered list above. Do not invent sources.

Be thorough and factual."""

    try:
        async with httpx.AsyncClient(timeout=90.0) as client:
            resp = await client.post(f"{OLLAMA}/api/chat", json={
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "options": {"temperature": 0.4, "num_predict": 8192, "num_ctx": 32768},
            })
            data = resp.json()
            return data.get("message", {}).get("content", "Unable to generate summary.")
    except Exception as exc:
        return f"# Research Summary: {topic}\n\nError generating summary: {exc}\n\n## Findings\n\n{findings}"


def _save_markdown_as_pdf(md_content: str, pdf_path: Path, title: str = "") -> None:
    """Convert Markdown content to a human-readable PDF and save to disk."""
    if not _WEASYPRINT_AVAILABLE or not _MARKDOWN_AVAILABLE:
        return
    html_body = _markdown_lib.markdown(md_content, extensions=["tables", "fenced_code", "toc", "nl2br"])
    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{title}</title>
<style>
@page {{ size: A4; margin: 2.5cm; }}
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Helvetica Neue', sans-serif; font-size: 11pt; line-height: 1.7; color: #1a1a1a; }}
h1 {{ font-size: 22pt; border-bottom: 2px solid #333; padding-bottom: 8px; margin-top: 0; margin-bottom: 16px; }}
h2 {{ font-size: 16pt; border-bottom: 1px solid #ccc; padding-bottom: 6px; margin-top: 28px; color: #2c3e50; }}
h3 {{ font-size: 13pt; margin-top: 20px; color: #34495e; }}
p {{ margin: 8px 0; text-align: justify; }}
ul, ol {{ margin: 8px 0; padding-left: 24px; }}
li {{ margin: 4px 0; }}
code {{ background: #f4f4f4; padding: 2px 5px; border-radius: 3px; font-size: 0.88em; font-family: 'SF Mono', 'Fira Code', Menlo, monospace; }}
pre {{ background: #f8f8f8; padding: 14px 18px; border-radius: 6px; overflow-x: auto; border: 1px solid #e8e8e8; font-size: 0.85em; line-height: 1.5; }}
pre code {{ background: none; padding: 0; }}
table {{ border-collapse: collapse; width: 100%; margin: 14px 0; font-size: 0.92em; }}
th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
th {{ background: #f0f0f0; font-weight: 600; }}
tr:nth-child(even) {{ background: #fafafa; }}
blockquote {{ border-left: 4px solid #3498db; margin: 14px 0; padding: 10px 18px; color: #555; background: #f9fbfd; border-radius: 0 4px 4px 0; }}
a {{ color: #2980b9; text-decoration: none; }}
hr {{ border: none; border-top: 1px solid #ddd; margin: 20px 0; }}
</style></head><body>{html_body}</body></html>"""
    doc = _weasyprint.HTML(string=html)
    doc.write_pdf(str(pdf_path))
    log.info("PDF saved: %s", pdf_path.name)


def _get_project_knowledge_context(project_id: str, max_chars: int = 8000) -> str:
    """Load project knowledge and format as context for the LLM."""
    try:
        knowledge_file = PROJECTS_DIR / project_id / "knowledge.json"
        if not knowledge_file.exists():
            return ""
        knowledge = json.loads(knowledge_file.read_text(encoding="utf-8"))

        findings = knowledge.get("findings", [])
        sources = knowledge.get("sources", [])

        if not findings and not sources:
            return ""

        parts = []
        parts.append("--- PROJECT KNOWLEDGE BASE ---")
        parts.append(f"Project: {knowledge.get('name', project_id)}")
        parts.append("")

        if findings:
            parts.append("Key Findings:")
            for i, f in enumerate(findings[-10:], 1):
                parts.append(f"{i}. [{f.get('topic', 'Unknown')}] {f.get('summary', '')[:300]}")
            parts.append("")

        if sources:
            parts.append("Available Sources (use [N] to cite inline):")
            for i, s in enumerate(sources[-15:], 1):
                parts.append(f"[{i}] {s.get('title', 'Unknown')}: {s.get('url', '')}")

        parts.append("---")

        context = "\n".join(parts)
        return context[:max_chars]
    except Exception:
        return ""


_tool_registry_cache: list[dict] | None = None
_tool_registry_mtime: float = 0


def _load_tool_registry() -> list[dict]:
    global _tool_registry_cache, _tool_registry_mtime
    try:
        mtime = _REGISTRY_FILE.stat().st_mtime if _REGISTRY_FILE.exists() else 0
        if _tool_registry_cache is not None and mtime == _tool_registry_mtime:
            return _tool_registry_cache
        data = json.loads(_REGISTRY_FILE.read_text(encoding="utf-8"))
        _tool_registry_cache = data
        _tool_registry_mtime = mtime
        return data
    except Exception:
        return []


def _save_tool_registry(registry: list[dict]) -> None:
    global _tool_registry_cache, _tool_registry_mtime
    _REGISTRY_FILE.write_text(json.dumps(registry, indent=2, ensure_ascii=False), encoding="utf-8")
    _tool_registry_cache = registry
    _tool_registry_mtime = _REGISTRY_FILE.stat().st_mtime

def _validate_tool_code(code: str) -> tuple[bool, str]:
    """Check tool code for blocked patterns."""
    for pattern in _TOOL_BLOCKED_PATTERNS:
        if pattern in code:
            return False, f"Blocked pattern found: {pattern}"
    return True, ""

def _execute_tool_sandboxed(tool_name: str, params: dict, timeout: float = 10.0) -> dict:
    """Execute a tool in an isolated subprocess for safety."""
    import sys as _sys

    registry = _load_tool_registry()
    tool_entry = next((t for t in registry if t["name"] == tool_name and t.get("enabled", True)), None)
    if not tool_entry:
        return {"error": f"Tool '{tool_name}' not found or disabled"}

    module_path = PLUGINS_DIR / tool_entry["module"]
    if not module_path.exists():
        return {"error": f"Tool module not found: {tool_entry['module']}"}

    expected_params = tool_entry.get("parameters", {})
    for param_name, param_spec in expected_params.items():
        if param_spec.get("required") and param_name not in params:
            return {"error": f"Missing required parameter: {param_name}"}

    # Build a temporary runner script that imports the tool and calls run()
    runner_code = (
        "import json, sys\n"
        "sys.path.insert(0, '')\n"
        "import importlib.util\n"
        f"spec = importlib.util.spec_from_file_location('tool', {str(module_path)!r})\n"
        "mod = importlib.util.module_from_spec(spec)\n"
        "spec.loader.exec_module(mod)\n"
        "if not hasattr(mod, 'run'):\n"
        "    print(json.dumps({'__sandbox_error': 'No run() function'}))\n"
        "    sys.exit(0)\n"
        f"params = json.loads({json.dumps(json.dumps(params))})\n"
        "try:\n"
        "    result = mod.run(**params)\n"
        "    print(json.dumps({'__sandbox_result': result}, default=str))\n"
        "except Exception as e:\n"
        "    print(json.dumps({'__sandbox_error': str(e)}))\n"
    )

    runner_file = None
    try:
        # Write runner to a temp file
        fd, runner_path = tempfile.mkstemp(suffix=".py", prefix="tool_runner_")
        runner_file = runner_path
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            f.write(runner_code)

        # Use the same Python interpreter (venv-aware)
        python_exe = _sys.executable

        # Build a minimal environment: inherit only essential vars
        safe_env = {
            "PATH": os.environ.get("PATH", ""),
            "HOME": os.environ.get("HOME", ""),
            "LANG": os.environ.get("LANG", "en_US.UTF-8"),
            "PYTHONPATH": "",
            "PYTHONDONTWRITEBYTECODE": "1",
        }
        # On macOS, DYLD_LIBRARY_PATH may be needed
        if "DYLD_LIBRARY_PATH" in os.environ:
            safe_env["DYLD_LIBRARY_PATH"] = os.environ["DYLD_LIBRARY_PATH"]

        proc = subprocess.run(
            [python_exe, runner_path],
            capture_output=True,
            text=True,
            timeout=timeout,
            env=safe_env,
            cwd=str(PLUGINS_DIR),
        )

        stdout = proc.stdout.strip()
        stderr = proc.stderr.strip()

        if proc.returncode != 0:
            error_msg = stderr or f"Process exited with code {proc.returncode}"
            _increment_tool_failure(tool_name)
            _log_tool_run(tool_name, params, None, error_msg)
            log.warning("Tool %s sandbox failed (exit %d): %s", tool_name, proc.returncode, error_msg)
            return {"error": f"Tool execution failed: {error_msg}"}

        if not stdout:
            _increment_tool_failure(tool_name)
            _log_tool_run(tool_name, params, None, "No output from sandbox")
            return {"error": "Tool produced no output"}

        # Parse the last line of stdout as JSON (tool may print debug info before)
        output_line = stdout.strip().split("\n")[-1]
        try:
            output = json.loads(output_line)
        except json.JSONDecodeError:
            _increment_tool_failure(tool_name)
            _log_tool_run(tool_name, params, None, f"Invalid JSON output: {output_line[:200]}")
            return {"error": f"Tool returned invalid output"}

        if "__sandbox_error" in output:
            _increment_tool_failure(tool_name)
            _log_tool_run(tool_name, params, None, output["__sandbox_error"])
            return {"error": f"Tool execution failed: {output['__sandbox_error']}"}

        result = output.get("__sandbox_result")

        for t in registry:
            if t["name"] == tool_name:
                t["usage_count"] = t.get("usage_count", 0) + 1
                t["last_used"] = datetime.now(timezone.utc).isoformat()
                t["consecutive_failures"] = 0
                break
        _save_tool_registry(registry)
        _log_tool_run(tool_name, params, result, None)
        log.info("Tool %s executed successfully (sandboxed)", tool_name)
        return {"result": result}

    except subprocess.TimeoutExpired:
        _increment_tool_failure(tool_name)
        _log_tool_run(tool_name, params, None, "Timeout")
        log.warning("Tool %s sandbox timed out after %ss", tool_name, timeout)
        return {"error": f"Tool '{tool_name}' timed out after {timeout}s"}
    except Exception as exc:
        _increment_tool_failure(tool_name)
        _log_tool_run(tool_name, params, None, str(exc))
        log.warning("Tool %s sandbox error: %s", tool_name, exc)
        return {"error": f"Sandbox execution failed: {exc}"}
    finally:
        if runner_file and os.path.exists(runner_file):
            os.unlink(runner_file)


def _execute_tool_legacy(tool_name: str, params: dict, timeout: float = 10.0) -> dict:
    """Load and execute a tool in-process (legacy fallback)."""
    import importlib.util
    import concurrent.futures

    registry = _load_tool_registry()
    tool_entry = next((t for t in registry if t["name"] == tool_name and t.get("enabled", True)), None)
    if not tool_entry:
        return {"error": f"Tool '{tool_name}' not found or disabled"}

    module_path = PLUGINS_DIR / tool_entry["module"]
    if not module_path.exists():
        return {"error": f"Tool module not found: {tool_entry['module']}"}

    expected_params = tool_entry.get("parameters", {})
    for param_name, param_spec in expected_params.items():
        if param_spec.get("required") and param_name not in params:
            return {"error": f"Missing required parameter: {param_name}"}

    try:
        spec = importlib.util.spec_from_file_location(f"tool_{tool_name}_{time.time()}", str(module_path))
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)

        if not hasattr(module, "run"):
            return {"error": f"Tool '{tool_name}' has no run() function"}

        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(module.run, **params)
            result = future.result(timeout=timeout)

        for t in registry:
            if t["name"] == tool_name:
                t["usage_count"] = t.get("usage_count", 0) + 1
                t["last_used"] = datetime.now(timezone.utc).isoformat()
                t["consecutive_failures"] = 0
                break
        _save_tool_registry(registry)
        _log_tool_run(tool_name, params, result, None)
        log.info("Tool %s executed successfully (legacy)", tool_name)
        return {"result": result}

    except concurrent.futures.TimeoutError:
        _increment_tool_failure(tool_name)
        _log_tool_run(tool_name, params, None, "Timeout")
        log.warning("Tool %s timed out after %ss", tool_name, timeout)
        return {"error": f"Tool '{tool_name}' timed out after {timeout}s"}
    except Exception as exc:
        _increment_tool_failure(tool_name)
        _log_tool_run(tool_name, params, None, str(exc))
        log.warning("Tool %s failed: %s", tool_name, exc)
        return {"error": f"Tool execution failed: {exc}"}


def _execute_tool(tool_name: str, params: dict, timeout: float = 10.0) -> dict:
    """Execute a tool, trying sandboxed subprocess first, falling back to legacy in-process."""
    result = _execute_tool_sandboxed(tool_name, params, timeout)
    if result.get("error", "").startswith("Sandbox execution failed:"):
        log.info("Sandbox failed for %s, falling back to legacy execution", tool_name)
        return _execute_tool_legacy(tool_name, params, timeout)
    return result

def _increment_tool_failure(tool_name: str) -> None:
    """Increment failure counter; auto-disable after 3 consecutive failures."""
    registry = _load_tool_registry()
    for t in registry:
        if t["name"] == tool_name:
            fails = t.get("consecutive_failures", 0) + 1
            t["consecutive_failures"] = fails
            if fails >= 3:
                t["enabled"] = False
                log.warning("Tool %s auto-disabled after %d consecutive failures", tool_name, fails)
            break
    _save_tool_registry(registry)

def _log_tool_run(name: str, params: dict, result, error: str | None) -> None:
    log_file = PLUGINS_DIR / "logs" / "tool_runs.json"
    try:
        logs = json.loads(log_file.read_text(encoding="utf-8")) if log_file.exists() else []
    except Exception:
        logs = []
    logs.append({
        "tool": name,
        "params": {k: str(v)[:100] for k, v in params.items()},
        "result": str(result)[:500] if result else None,
        "error": error,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    })
    logs = logs[-200:]
    try:
        log_file.write_text(json.dumps(logs, indent=2, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass

def _get_tools_summary() -> str:
    """Format available tools as a string for the system prompt."""
    registry = _load_tool_registry()
    enabled = [t for t in registry if t.get("enabled", True)]
    lines = []
    lines.append("--- AVAILABLE TOOLS ---")
    lines.append("CRITICAL: You have a tool execution system. To call a tool, you MUST use this EXACT syntax:")
    lines.append("<<TOOL:tool_name(param1=value1, param2=value2)>>")
    lines.append("")
    lines.append("RULES:")
    lines.append("- Use ONLY the <<TOOL:...>> syntax. Do NOT use any other tool-calling format.")
    lines.append("- Do NOT use function calling, <|tool_calls|>, or any model-native tool syntax.")
    lines.append("- Do NOT hallucinate or make up data. If you need real data, use a tool.")
    lines.append("- The system will execute the tool and return real results to you.")
    lines.append("")
    if enabled:
        lines.append("Available tools:")
        for t in enabled:
            params_str = ", ".join(
                f"{k}={v.get('type', 'string')}" + (" REQUIRED" if v.get("required") else "")
                for k, v in (t.get("parameters") or {}).items()
            )
            lines.append(f"  • {t['name']}({params_str}) — {t.get('description', '')}")
        lines.append("")
        lines.append("Example: <<TOOL:weather(city=London)>>")
    else:
        lines.append("No tools are currently installed.")
    lines.append("")
    lines.append("If you need external data and no tool above covers it, respond with:")
    lines.append("<<BUILD_TOOL:description of what capability you need>>")
    lines.append("The system will automatically research, build, and register a new tool.")
    lines.append("---")
    return "\n".join(lines)

def _load_memories() -> list[str]:
    try:
        data = json.loads(_MEMORY_FILE.read_text(encoding="utf-8"))
        return [str(m) for m in data if isinstance(m, str)]
    except Exception:
        return []

def _save_memories(memories: list[str]) -> None:
    _MEMORY_FILE.write_text(json.dumps(memories, indent=2, ensure_ascii=False), encoding="utf-8")

def _add_memory(text: str) -> None:
    memories = _load_memories()
    text = text.strip()
    if text and text not in memories:
        memories.append(text)
        _save_memories(memories)
        log.info("Memory added: %s", text[:80])

def _remove_memory(index: int) -> bool:
    memories = _load_memories()
    if 0 <= index < len(memories):
        removed = memories.pop(index)
        _save_memories(memories)
        log.info("Memory removed: %s", removed[:80])
        return True
    return False

def _get_memory_context() -> str:
    memories = _load_memories()
    if not memories:
        return ""
    lines = ["--- USER PREFERENCES (always follow these) ---"]
    for m in memories:
        lines.append(f"- {m}")
    lines.append("---")
    return "\n".join(lines)

_TOOL_CALL_RE = re.compile(r'<<TOOL:(\w+)\(([^>]*?)\)>>')
_BUILD_TOOL_RE = re.compile(r'<<BUILD_TOOL:(.+?)>>')
_ALL_TOOL_TAGS_RE = re.compile(r'<<(?:TOOL:\w+\([^>]*?\)|BUILD_TOOL:.+?)>>')


def _parse_tool_calls(text: str) -> list[tuple[str, dict]]:
    """Parse <<TOOL:name(key=value, key2=value2)>> patterns. Handles quoted values with commas."""
    calls = []
    for match in _TOOL_CALL_RE.finditer(text):
        name = match.group(1)
        params_str = match.group(2).strip()
        params = {}
        if params_str:
            parts = re.split(r',\s*(?=\w+=)', params_str)
            for part in parts:
                if "=" in part:
                    k, v = part.split("=", 1)
                    params[k.strip()] = v.strip().strip("'\"")
        calls.append((name, params))
    return calls


def _strip_tags_for_display(text: str) -> str:
    """Remove <<TOOL:...>>, <<BUILD_TOOL:...>> tags, and native model tool syntax from text."""
    cleaned = _strip_native_tool_syntax(text)
    cleaned = _ALL_TOOL_TAGS_RE.sub('', cleaned)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned).strip()
    return cleaned


_NATIVE_TOOL_PATTERNS = [
    re.compile(r'<｜tool▁calls▁begin｜>.*?<｜tool▁calls▁end｜>', re.DOTALL),
    re.compile(r'<｜tool▁outputs▁begin｜>.*?<｜tool▁outputs▁end｜>', re.DOTALL),
    re.compile(r'<\|tool_calls?\|>.*?(?:<\|/tool_calls?\|>|$)', re.DOTALL),
    re.compile(r'<\|tool_outputs?\|>.*?(?:<\|/tool_outputs?\|>|$)', re.DOTALL),
    re.compile(r'<tool_call>.*?</tool_call>', re.DOTALL),
    re.compile(r'<function_call>.*?</function_call>', re.DOTALL),
    re.compile(r'```tool_code.*?```', re.DOTALL),
]


def _strip_native_tool_syntax(text: str) -> str:
    """Remove native model tool-calling markup from response text."""
    cleaned = text
    for pattern in _NATIVE_TOOL_PATTERNS:
        cleaned = pattern.sub('', cleaned)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned).strip()
    return cleaned


def _pre_match_tools(user_message: str) -> list[tuple[str, dict]]:
    """Match the user's message against installed tools using keyword analysis.
    
    Returns a list of (tool_name, params) to execute BEFORE sending to the LLM.
    This is the fast path — no LLM call needed, works with any model.
    """
    if not user_message:
        return []
    
    registry = _load_tool_registry()
    enabled = [t for t in registry if t.get("enabled", True)]
    if not enabled:
        return []
    
    msg = user_message.lower().strip()
    matches = []
    
    for tool in enabled:
        name = tool["name"]
        desc = (tool.get("description") or "").lower()
        params = tool.get("parameters", {})
        
        name_words = set(name.replace("_", " ").split())
        desc_words = set(re.findall(r'\b[a-z]{3,}\b', desc))
        tool_keywords = name_words | desc_words
        tool_keywords -= {"get", "current", "any", "for", "the", "and", "from", "with", "using", "conditions"}
        
        msg_words = set(re.findall(r'\b[a-z]{3,}\b', msg))
        overlap = tool_keywords & msg_words
        
        if not overlap:
            continue
        
        extracted_params = {}
        for param_name, param_spec in params.items():
            param_desc = (param_spec.get("description") or "").lower()
            
            if "city" in param_name.lower() or "city" in param_desc or "location" in param_name.lower():
                city_patterns = [
                    r'\b(?:in|for|at|of)\s+([A-Z][a-zA-Z\s]+?)(?:\s+(?:today|now|right now|currently|tonight|tomorrow))?\s*[?.!,]?\s*$',
                    r'\b(?:in|for|at|of)\s+([A-Z][a-zA-Z\s]+?)(?:\s*[?.!,]|$)',
                    r'\b(?:weather|temperature|forecast|climate)\s+(?:in\s+|for\s+|at\s+)?([A-Z][a-zA-Z\s]+?)(?:\s*[?.!,]|$)',
                    r'([A-Z][a-zA-Z\s]+?)\s+(?:weather|temperature|forecast|climate)',
                ]
                for pattern in city_patterns:
                    match = re.search(pattern, user_message)
                    if match:
                        extracted_params[param_name] = match.group(1).strip()
                        break
            
            elif "currency" in param_name.lower() or "currency" in param_desc:
                cur_match = re.search(r'\b([A-Z]{3})\b', user_message)
                if cur_match:
                    extracted_params[param_name] = cur_match.group(1)
            
            elif "symbol" in param_name.lower() or "ticker" in param_name.lower():
                ticker_match = re.search(r'\b([A-Z]{1,5})\b', user_message)
                if ticker_match:
                    extracted_params[param_name] = ticker_match.group(1)
            
            elif param_spec.get("type") == "number":
                num_match = re.search(r'\b(\d+(?:\.\d+)?)\b', user_message)
                if num_match:
                    extracted_params[param_name] = num_match.group(1)
        
        required_params = {k for k, v in params.items() if v.get("required")}
        if required_params and not required_params.issubset(extracted_params.keys()):
            continue
        
        if len(overlap) >= 1:
            matches.append((name, extracted_params))
    
    return matches


async def _build_tool(description: str, model: str) -> dict:
    """Autonomously research, build, test, and register a new tool."""
    log.info("Building tool: %s", description)
    queries = await _generate_search_queries(f"free API no key required for: {description}", 3, model)
    api_info = []
    for q in queries:
        results = await _do_web_search(q, max_results=3)
        for r in results[:2]:
            page = await _fetch_page_content(r.get("href", ""), max_chars=3000)
            if page:
                api_info.append(page)

    combined_research = "\n\n---\n\n".join(api_info[:5])[:15000]

    tool_prompt = f"""Write a Python tool module for: "{description}"

Research findings about available APIs:
{combined_research}

The module MUST follow this EXACT format:
TOOL_NAME = "short_name"  # lowercase, underscores, GENERIC name (e.g. "weather" not "weather_london")
TOOL_DESCRIPTION = "What this tool does in one sentence"
TOOL_PARAMETERS = {{
    "param_name": {{"type": "string", "description": "What this param is", "required": True}}
}}

def run(**kwargs) -> dict:
    import httpx
    # implementation
    return {{"key": "value"}}

CRITICAL RULES:
- The tool MUST be GENERIC, not specific to one city/item/thing. Use parameters for the variable parts.
- TOOL_NAME must be short and generic: "weather", "stock_price", "currency", NOT "weather_cape_town"
- Use ONLY free APIs that require NO API keys. Strongly prefer: Open-Meteo (weather), exchangerate.host (currency), etc.
- Do NOT use any API that needs a key, token, or registration. If you put "YOUR_API_KEY" anywhere, the tool is BROKEN.
- Only import from: httpx, json, datetime, re, math, urllib, html, csv, collections, time, calendar, decimal, statistics, base64, hashlib
- The run() function MUST accept **kwargs and return a dict
- Handle errors gracefully with try/except, return {{"error": "message"}} on failure
- Do NOT use os, subprocess, eval, exec, open(), pathlib, shutil, glob
- Keep it simple and focused on one task

For weather specifically, use Open-Meteo API:
  https://api.open-meteo.com/v1/forecast?latitude=XX&longitude=YY&current=temperature_2m,wind_speed_10m,relative_humidity_2m,weather_code
  Use https://geocoding-api.open-meteo.com/v1/search?name=CITY&count=1 to get lat/lon from city name.

Return ONLY the Python code, no markdown fences, no explanations."""

    fix_prompt = tool_prompt
    for attempt in range(3):
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                resp = await client.post(f"{OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": fix_prompt}],
                    "stream": False,
                    "options": {"temperature": 0.2, "num_predict": 4096, "num_ctx": 32768},
                })
                data = resp.json()
                code = data.get("message", {}).get("content", "")

            code = code.strip()
            if code.startswith("```"):
                lines = code.split("\n")
                if lines[-1].strip() == "```":
                    lines = lines[1:-1]
                else:
                    lines = lines[1:]
                code = "\n".join(lines)

            valid, error = _validate_tool_code(code)
            if not valid:
                fix_prompt = f"The previous code had a security issue: {error}. Rewrite the entire tool module without using blocked patterns. Return only Python code."
                continue

            import importlib.util
            temp_path = PLUGINS_DIR / "tools" / "_temp_build.py"
            temp_path.write_text(code, encoding="utf-8")

            spec = importlib.util.spec_from_file_location(f"_temp_build_{time.time()}", str(temp_path))
            module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(module)

            tool_name = getattr(module, "TOOL_NAME", None)
            tool_desc = getattr(module, "TOOL_DESCRIPTION", None)
            tool_params = getattr(module, "TOOL_PARAMETERS", {})
            run_fn = getattr(module, "run", None)

            if not tool_name or not run_fn:
                fix_prompt = f"The module is missing TOOL_NAME or run(). Rewrite the complete module. Return only Python code."
                temp_path.unlink(missing_ok=True)
                continue

            try:
                import concurrent.futures
                test_params = {}
                for k, v in tool_params.items():
                    if v.get("type") == "number":
                        test_params[k] = 0
                    else:
                        test_params[k] = "London" if "city" in k.lower() else "test"
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                    future = executor.submit(run_fn, **test_params)
                    test_result = future.result(timeout=10.0)
            except Exception as test_err:
                fix_prompt = f"The tool threw an error during testing with params {test_params}: {test_err}. Fix the run() function. Return only Python code."
                temp_path.unlink(missing_ok=True)
                continue

            final_path = PLUGINS_DIR / "tools" / f"{tool_name}.py"
            temp_path.rename(final_path)

            registry = _load_tool_registry()
            registry = [t for t in registry if t["name"] != tool_name]
            registry.append({
                "name": tool_name,
                "description": tool_desc or description,
                "parameters": tool_params,
                "module": f"tools/{tool_name}.py",
                "created": datetime.now(timezone.utc).isoformat(),
                "usage_count": 0,
                "last_used": None,
                "enabled": True,
                "consecutive_failures": 0,
            })
            _save_tool_registry(registry)
            log.info("Tool %s built successfully", tool_name)
            return {"ok": True, "name": tool_name, "description": tool_desc}

        except Exception as exc:
            fix_prompt = f"Failed to load the module: {exc}. Rewrite the complete tool module. Return only Python code."
            temp_path = PLUGINS_DIR / "tools" / "_temp_build.py"
            temp_path.unlink(missing_ok=True)
            continue

    log.warning("Tool build failed after 3 attempts: %s", description)
    return {"ok": False, "error": "Failed to build tool after 3 attempts"}


async def _build_tool_preview(description: str, model: str) -> dict:
    """Generate tool code for preview without registering or testing it."""
    log.info("Building tool preview: %s", description)
    queries = await _generate_search_queries(f"free API no key required for: {description}", 3, model)
    api_info = []
    for q in queries:
        results = await _do_web_search(q, max_results=3)
        for r in results[:2]:
            page = await _fetch_page_content(r.get("href", ""), max_chars=3000)
            if page:
                api_info.append(page)

    combined_research = "\n\n---\n\n".join(api_info[:5])[:15000]

    tool_prompt = f"""Write a Python tool module for: "{description}"

Research findings about available APIs:
{combined_research}

The module MUST follow this EXACT format:
TOOL_NAME = "short_name"
TOOL_DESCRIPTION = "What this tool does in one sentence"
TOOL_PARAMETERS = {{
    "param_name": {{"type": "string", "description": "What this param is", "required": True}}
}}

def run(**kwargs) -> dict:
    import httpx
    # implementation
    return {{"key": "value"}}

CRITICAL RULES:
- The tool MUST be GENERIC, not specific to one city/item/thing. Use parameters for the variable parts.
- TOOL_NAME must be short and generic: "weather", "stock_price", "currency"
- Use ONLY free APIs that require NO API keys.
- Only import from: httpx, json, datetime, re, math, urllib, html, csv, collections, time, calendar, decimal, statistics, base64, hashlib
- The run() function MUST accept **kwargs and return a dict
- Handle errors gracefully with try/except, return {{"error": "message"}} on failure
- Do NOT use os, subprocess, eval, exec, open(), pathlib, shutil, glob

Return ONLY the Python code, no markdown fences, no explanations."""

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(f"{OLLAMA}/api/chat", json={
                "model": model,
                "messages": [{"role": "user", "content": tool_prompt}],
                "stream": False,
                "options": {"temperature": 0.2, "num_predict": 4096, "num_ctx": 32768},
            })
            data = resp.json()
            code = data.get("message", {}).get("content", "")

        code = code.strip()
        if code.startswith("```"):
            lines = code.split("\n")
            if lines[-1].strip() == "```":
                lines = lines[1:-1]
            else:
                lines = lines[1:]
            code = "\n".join(lines)

        valid, error = _validate_tool_code(code)
        if not valid:
            return {"ok": False, "error": f"Generated code has security issue: {error}", "code": code}

        # Extract metadata from generated code without executing
        name_match = re.search(r'TOOL_NAME\s*=\s*["\']([^"\']+)["\']', code)
        desc_match = re.search(r'TOOL_DESCRIPTION\s*=\s*["\']([^"\']+)["\']', code)
        tool_name = name_match.group(1) if name_match else "unknown"
        tool_desc = desc_match.group(1) if desc_match else description

        return {
            "ok": True,
            "preview": True,
            "name": tool_name,
            "description": tool_desc,
            "code": code,
        }

    except Exception as exc:
        log.warning("Tool preview build failed: %s", exc)
        return {"ok": False, "error": f"Failed to generate tool code: {exc}"}


async def stream_ollama_response(path: str, body: dict, *, write_timeout: float):
    try:
        async with httpx.AsyncClient(
            timeout=httpx.Timeout(connect=5.0, read=None, write=write_timeout, pool=5.0)
        ) as client:
            async with client.stream("POST", f"{OLLAMA}{path}", json=body) as resp:
                if resp.status_code >= 400:
                    detail = (await resp.aread()).decode("utf-8", errors="replace").strip()
                    reason = detail or resp.reason_phrase
                    yield _ndjson_error(f"Ollama returned {resp.status_code}: {reason}")
                    return
                async for chunk in resp.aiter_bytes():
                    if chunk:
                        yield chunk
    except httpx.ConnectError:
        yield _ndjson_error("Cannot connect to Ollama. Start it with: ollama serve")
    except httpx.HTTPError as exc:
        yield _ndjson_error(f"Ollama request failed: {exc}")
    except Exception as exc:
        yield _ndjson_error(str(exc))


def _token_table_lines(active: str | None, prompt_req: int, completion_req: int) -> list[str]:
    stats    = _token_stats
    name_col = max(15, max((len(k) for k in stats), default=0) + 4)
    total_p  = sum(v[0] for v in stats.values())
    total_c  = sum(v[1] for v in stats.values())
    grand    = total_p + total_c
    W = name_col + 39
    C = "─"
    if active is not None:
        head = f" Token Usage  ·  +{prompt_req + completion_req:,} this request "
    else:
        head = " Session Token Summary  ·  Server Shutdown "
    head = head.center(W)
    lines = [
        f"┌{C * W}┐",
        f"│{head}│",
        f"├{C * name_col}┬{C * 12}┬{C * 12}┬{C * 12}┤",
        f"│ {'Client':<{name_col - 2}} │ {'Prompt':>10} │ {'Completion':>10} │ {'Total':>10} │",
        f"├{C * name_col}┼{C * 12}┼{C * 12}┼{C * 12}┤",
    ]
    for name, (p, c) in sorted(stats.items(), key=lambda x: -(x[1][0] + x[1][1])):
        marker = " ◀" if name == active else ""
        lines.append(
            f"│ {(name + marker):<{name_col - 2}} │ {p:>10,} │ {c:>10,} │ {p + c:>10,} │"
        )
    lines += [
        f"├{C * name_col}┼{C * 12}┼{C * 12}┼{C * 12}┤",
        f"│ {'TOTAL':<{name_col - 2}} │ {total_p:>10,} │ {total_c:>10,} │ {grand:>10,} │",
        f"└{C * name_col}┴{C * 12}┴{C * 12}┴{C * 12}┘",
    ]
    return lines


def _print_token_table(display_name: str, prompt_req: int, completion_req: int) -> None:
    log.info("\n" + "\n".join(_token_table_lines(display_name, prompt_req, completion_req)))


def _print_shutdown_summary() -> None:
    if not _token_stats:
        return
    log.info("\n" + "\n".join(_token_table_lines(None, 0, 0)))


atexit.register(_print_shutdown_summary)


def _tally_done_line(line: bytes, display_name: str) -> None:
    line = line.strip()
    if not line:
        return
    try:
        data = json.loads(line)
        if data.get("done"):
            prompt_req     = data.get("prompt_eval_count", 0)
            completion_req = data.get("eval_count", 0)
            entry = _token_stats.setdefault(display_name, [0, 0])
            entry[0] += prompt_req
            entry[1] += completion_req
            _save_token_stats()
            _print_token_table(display_name, prompt_req, completion_req)
    except (json.JSONDecodeError, AttributeError):
        pass


async def _chat_stream_with_token_log(body: dict, client_host: str):
    display_name = body.get("user", "").strip() or client_host
    ollama_body  = {k: v for k, v in body.items() if k != "user"}
    buf = b""
    async for chunk in stream_ollama_response("/api/chat", ollama_body, write_timeout=120.0):
        buf += chunk
        while b"\n" in buf:
            line, buf = buf.split(b"\n", 1)
            _tally_done_line(line, display_name)
        yield chunk
    if buf.strip():
        _tally_done_line(buf, display_name)


async def _chat_with_tool_execution(body: dict, client_host: str):
    """Wrap chat streaming with automatic tool execution and tool building.
    
    Flow:
    1. Buffer the model's full response (don't stream raw tool tags to client)
    2. Strip native tool syntax and display clean text
    3. If <<BUILD_TOOL:...>> found: build the tool, then re-run
    4. If <<TOOL:...>> found: execute tools, send results back to model, stream final answer
    5. If neither: just stream the clean response
    """
    full_response = ""

    async for chunk in _chat_stream_with_token_log(body, client_host):
        for line in chunk.decode("utf-8", errors="replace").split("\n"):
            line = line.strip()
            if not line:
                continue
            try:
                data = json.loads(line)
                content = data.get("message", {}).get("content", "")
                if content:
                    full_response += content
            except (json.JSONDecodeError, AttributeError):
                pass
        yield chunk

    cleaned_response = _strip_tags_for_display(full_response)

    build_requests = _BUILD_TOOL_RE.findall(full_response)
    if build_requests:
        for desc in build_requests:
            msg = f"\n\n---\n\U0001f528 *Building new tool: {desc.strip()}...*\n"
            yield (json.dumps({"message": {"role": "assistant", "content": msg}, "done": False}) + "\n").encode()

            build_result = await _build_tool(desc.strip(), body.get("model", FALLBACK_MODEL))

            if build_result.get("ok"):
                status = f"\u2714 Tool '{build_result['name']}' created and ready to use!\n\n"
            else:
                status = f"\u26a0\ufe0f Could not build tool: {build_result.get('error', 'Unknown error')}\n\n"
            yield (json.dumps({"message": {"role": "assistant", "content": status}, "done": False}) + "\n").encode()

        follow_up = dict(body)
        msgs = list(follow_up.get("messages", []))
        msgs.append({"role": "assistant", "content": cleaned_response})
        msgs.append({"role": "user", "content": "The tools have been built. Now answer the original question using the newly available tools. Use <<TOOL:name(params)>> to call them."})
        follow_up["messages"] = msgs
        tools_summary = _get_tools_summary()
        for m in follow_up["messages"]:
            if m.get("role") == "system":
                m["content"] = m["content"] + "\n\n" + tools_summary
                break

        full_response = ""
        async for chunk in _chat_stream_with_token_log(follow_up, client_host):
            for line in chunk.decode("utf-8", errors="replace").split("\n"):
                line = line.strip()
                if not line:
                    continue
                try:
                    data = json.loads(line)
                    content = data.get("message", {}).get("content", "")
                    if content:
                        full_response += content
                except (json.JSONDecodeError, AttributeError):
                    pass
            yield chunk

    tool_calls = _parse_tool_calls(full_response)
    if not tool_calls:
        return

    for t_name, t_params in tool_calls:
        params_display = ", ".join(f"{k}={v}" for k, v in t_params.items())
        progress_msg = f"\n\U0001f527 *Running tool: {t_name}({params_display})...*\n"
        yield (json.dumps({"message": {"role": "assistant", "content": progress_msg}, "done": False}) + "\n").encode()

    tool_results = []
    for t_name, t_params in tool_calls:
        result = await asyncio.to_thread(_execute_tool, t_name, t_params)
        tool_results.append({"tool": t_name, "params": t_params, "result": result})
        result_preview = json.dumps(result.get("result", result), indent=2)[:800]
        yield (json.dumps({"message": {"role": "assistant", "content": f"\u2714 {t_name}: {result_preview}\n"}, "done": False}) + "\n").encode()

    results_text = "\n\n".join(
        f"Tool '{tr['tool']}' result: {json.dumps(tr['result'])}" for tr in tool_results
    )

    yield (json.dumps({"message": {"role": "assistant", "content": "\n---\n*Processing results...*\n\n"}, "done": False}) + "\n").encode()

    yield (json.dumps({"message": {"role": "assistant", "content": ""}, "done": True}) + "\n").encode()

    yield b"\n"

    follow_up = dict(body)
    msgs = list(follow_up.get("messages", []))
    msgs.append({"role": "assistant", "content": _strip_tags_for_display(full_response)})
    msgs.append({"role": "user", "content": f"Here are the tool execution results:\n{results_text}\n\nUse these results to provide your final answer. Do NOT include <<TOOL:...>> or <<BUILD_TOOL:...>> tags \u2014 the tools have already been executed. Just give a natural response using the data."})
    follow_up["messages"] = msgs

    async for chunk in _chat_stream_with_token_log(follow_up, client_host):
        yield chunk


@app.post("/api/chat")
async def chat(request: Request):
    body = await request.json()
    client_host = request.client.host if request.client else "unknown"
    log.debug("Chat request from %s model=%s messages=%d", client_host, body.get("model", "?"), len(body.get("messages", [])))
    memory_context = _get_memory_context()
    if memory_context:
        messages = body.get("messages", [])
        has_system = any(m.get("role") == "system" for m in messages)
        if has_system:
            for m in messages:
                if m.get("role") == "system":
                    m["content"] = memory_context + "\n\n" + m["content"]
                    break
        else:
            messages.insert(0, {"role": "system", "content": memory_context})
        body["messages"] = messages

    project_id = body.pop("project_id", None)
    if project_id:
        knowledge_context = _get_project_knowledge_context(project_id)
        if knowledge_context:
            messages = body.get("messages", [])
            has_system = any(m.get("role") == "system" for m in messages)
            if has_system:
                for m in messages:
                    if m.get("role") == "system":
                        m["content"] = m["content"] + "\n\n" + knowledge_context
                        break
            else:
                messages.insert(0, {"role": "system", "content": knowledge_context})
            body["messages"] = messages

    tools_summary = _get_tools_summary()
    if tools_summary:
        messages = body.get("messages", [])
        has_system = any(m.get("role") == "system" for m in messages)
        if has_system:
            for m in messages:
                if m.get("role") == "system":
                    m["content"] = m["content"] + "\n\n" + tools_summary
                    break
        else:
            messages.insert(0, {"role": "system", "content": tools_summary})
        body["messages"] = messages

    user_messages = [m for m in body.get("messages", []) if m.get("role") == "user"]
    last_user_msg = user_messages[-1]["content"] if user_messages else ""
    pre_matches = _pre_match_tools(last_user_msg)
    if pre_matches: log.info("Pre-matched tools: %s", [(n, p) for n, p in pre_matches])

    if pre_matches:
        tool_data_parts = []
        for t_name, t_params in pre_matches:
            result = await asyncio.to_thread(_execute_tool, t_name, t_params)
            if result.get("result"):
                tool_data_parts.append(
                    f"[Tool data — {t_name}({', '.join(f'{k}={v}' for k, v in t_params.items())})]\n"
                    f"{json.dumps(result['result'], indent=2)}"
                )

        if tool_data_parts:
            tool_data_block = (
                "\n\n--- LIVE TOOL DATA (real-time, use this to answer) ---\n"
                + "\n\n".join(tool_data_parts)
                + "\n--- END TOOL DATA ---"
            )
            messages = body.get("messages", [])
            for m in messages:
                if m.get("role") == "system":
                    m["content"] = m["content"] + tool_data_block
                    break
            else:
                messages.insert(0, {"role": "system", "content": tool_data_block.strip()})
            body["messages"] = messages

    return StreamingResponse(
        _chat_with_tool_execution(body, client_host),
        media_type="application/x-ndjson",
    )


@app.get("/api/tokens")
async def get_tokens():
    return JSONResponse(_token_stats)


@app.delete("/api/tokens")
async def reset_user_tokens(user: str = ""):
    if user and user in _token_stats:
        _token_stats[user] = [0, 0]
        _save_token_stats()
    return JSONResponse({"ok": True})


@app.post("/api/pull")
async def pull_model(request: Request):
    body = await request.json()
    return StreamingResponse(
        stream_ollama_response("/api/pull", body, write_timeout=30.0),
        media_type="application/x-ndjson",
    )


@app.post("/api/generate-image")
async def generate_image(request: Request):
    """Generate an image using a local Diffusers pipeline.
    Streams NDJSON progress events to match the frontend contract."""
    body = _apply_image_generation_caps(await request.json())

    async def _stream():
        try:
            # Check availability first
            status = _image_gen.get_status()
            if not status["available"]:
                yield json.dumps({"error": status["error"]}) + "\n"
                return

            # Signal model loading if not yet loaded
            if not status["model_loaded"]:
                yield json.dumps({"status": "Loading image model…"}) + "\n"

            yield json.dumps({"status": "Generating image…", "progress": 10}) + "\n"

            b64_image = await _image_gen.generate_image(
                prompt=body.get("prompt", ""),
                width=body.get("width", 768),
                height=body.get("height", 768),
                steps=body.get("steps", 10),
                negative_prompt=body.get("negative_prompt"),
                seed=body.get("seed"),
            )

            yield json.dumps({"status": "Generating image…", "progress": 90}) + "\n"
            yield json.dumps({"image": b64_image, "response": b64_image, "done": True}) + "\n"

        except Exception as exc:
            log.error("Image generation failed: %s", exc)
            yield json.dumps({"error": str(exc)}) + "\n"

    return StreamingResponse(_stream(), media_type="application/x-ndjson")


@app.post("/api/fetch-page")
async def fetch_page(request: Request):
    """Fetch a URL, strip HTML, return clean text."""
    if not _BS4_AVAILABLE:
        return JSONResponse(
            {"error": "beautifulsoup4 is required. Install with: pip install beautifulsoup4 lxml"},
            status_code=501,
        )
    body = await request.json()
    url = (body.get("url") or "").strip()
    if not url:
        return JSONResponse({"error": "No URL provided"}, status_code=400)
    max_chars = int(body.get("max_chars", 8000))

    try:
        async with httpx.AsyncClient(timeout=10.0, follow_redirects=True) as client:
            resp = await client.get(url)
            resp.raise_for_status()
    except Exception as exc:
        return JSONResponse({"error": f"Failed to fetch URL: {exc}"}, status_code=502)

    soup = BeautifulSoup(resp.text, "lxml")

    for tag in soup.find_all(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    title = soup.title.get_text(strip=True) if soup.title else ""
    content = soup.get_text(separator='\n', strip=True)
    content = content[:max_chars]

    return {"url": url, "title": title, "content": content}


PROJECTS_DIR = Path.home() / "OfflineAI-Projects"
PROJECTS_DIR.mkdir(parents=True, exist_ok=True)

PLUGINS_DIR = Path.home() / "OfflineAI-Plugins"
PLUGINS_DIR.mkdir(parents=True, exist_ok=True)
(PLUGINS_DIR / "tools").mkdir(exist_ok=True)
(PLUGINS_DIR / "logs").mkdir(exist_ok=True)
_REGISTRY_FILE = PLUGINS_DIR / "registry.json"
if not _REGISTRY_FILE.exists():
    _REGISTRY_FILE.write_text("[]", encoding="utf-8")

MEMORY_DIR = Path.home() / "OfflineAI-Memory"
MEMORY_DIR.mkdir(parents=True, exist_ok=True)
_MEMORY_FILE = MEMORY_DIR / "preferences.json"
if not _MEMORY_FILE.exists():
    _MEMORY_FILE.write_text("[]", encoding="utf-8")

_TOOL_BLOCKED_PATTERNS = [
    "os.system", "subprocess", "eval(", "exec(", "__import__",
    "open(", "pathlib", "shutil", "glob",
    "compile(",
    # Network access beyond httpx
    "socket", "requests", "urllib.request", "http.client",
    # Deserialization attacks
    "pickle", "marshal",
    # Native code access
    "ctypes", "cffi",
    # System manipulation
    "sys.exit", "os.environ",
    # Dynamic imports
    "importlib",
]


def _slugify(text: str) -> str:
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii')
    text = re.sub(r'[^\w\s-]', '', text.lower())
    return re.sub(r'[-\s]+', '-', text).strip('-')


def _resolve_project_path(project_id: str, *parts: str) -> Path | None:
    """Resolve a path within a project directory, returning None if it escapes."""
    project_dir = PROJECTS_DIR / project_id
    if parts:
        target = (project_dir / Path(*parts)).resolve()
    else:
        target = project_dir.resolve()
    if not str(target).startswith(str(project_dir.resolve())):
        return None
    return target


def _count_files(project_dir: Path) -> int:
    """Count all files in a project directory, excluding knowledge.json."""
    count = 0
    for f in project_dir.rglob("*"):
        if f.is_file() and f.name != "knowledge.json":
            count += 1
    return count


@app.get("/api/projects")
async def list_projects():
    """List all projects."""
    projects = []
    if PROJECTS_DIR.is_dir():
        for entry in sorted(PROJECTS_DIR.iterdir()):
            if entry.is_dir():
                knowledge_file = entry / "knowledge.json"
                if knowledge_file.is_file():
                    try:
                        data = json.loads(knowledge_file.read_text(encoding="utf-8"))
                        projects.append({
                            "id": entry.name,
                            "name": data.get("name", entry.name),
                            "description": data.get("description", ""),
                            "created": data.get("created", ""),
                            "sources_count": len(data.get("sources", [])),
                            "findings_count": len(data.get("findings", [])),
                            "files_count": _count_files(entry),
                        })
                    except (json.JSONDecodeError, OSError):
                        continue
    return {"projects": projects}


@app.post("/api/projects")
async def create_project(request: Request):
    """Create a new project."""
    body = await request.json()
    name = (body.get("name") or "").strip()
    description = (body.get("description") or "").strip()
    if not name:
        return JSONResponse({"error": "Project name is required"}, status_code=400)

    slug = _slugify(name)
    if not slug:
        return JSONResponse({"error": "Invalid project name"}, status_code=400)

    project_dir = PROJECTS_DIR / slug
    if project_dir.exists():
        return JSONResponse({"error": f"Project '{slug}' already exists"}, status_code=409)

    project_dir.mkdir(parents=True)
    (project_dir / "notes").mkdir()
    (project_dir / "sources").mkdir()
    (project_dir / "output").mkdir()

    created = datetime.now(timezone.utc).isoformat()
    knowledge = {
        "name": name,
        "description": description,
        "created": created,
        "sources": [],
        "findings": [],
    }
    (project_dir / "knowledge.json").write_text(json.dumps(knowledge, indent=2), encoding="utf-8")

    log.info("Project created: %s (%s)", name, slug)
    return {
        "id": slug,
        "name": name,
        "description": description,
        "created": created,
        "sources_count": 0,
        "findings_count": 0,
        "files_count": 0,
    }


@app.get("/api/projects/{project_id}")
async def get_project(project_id: str):
    """Get a single project's metadata."""
    project_dir = _resolve_project_path(project_id)
    if project_dir is None or not project_dir.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    knowledge_file = project_dir / "knowledge.json"
    if not knowledge_file.is_file():
        return JSONResponse({"error": "Project metadata missing"}, status_code=404)

    try:
        data = json.loads(knowledge_file.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as exc:
        return JSONResponse({"error": f"Failed to read project: {exc}"}, status_code=500)

    return {
        "id": project_id,
        "name": data.get("name", project_id),
        "description": data.get("description", ""),
        "created": data.get("created", ""),
        "sources_count": len(data.get("sources", [])),
        "findings_count": len(data.get("findings", [])),
        "files_count": _count_files(project_dir),
    }


@app.delete("/api/projects/{project_id}")
async def delete_project(project_id: str):
    """Delete an entire project."""
    project_dir = _resolve_project_path(project_id)
    if project_dir is None or not project_dir.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    shutil.rmtree(project_dir)
    log.info("Project deleted: %s", project_id)
    return {"ok": True}


@app.get("/api/projects/{project_id}/files")
async def list_project_files(project_id: str):
    """List all files in a project."""
    project_dir = _resolve_project_path(project_id)
    if project_dir is None or not project_dir.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    files = []
    for f in sorted(project_dir.rglob("*")):
        if f.is_file() and f.name != "knowledge.json":
            rel = f.relative_to(project_dir)
            stat = f.stat()
            files.append({
                "path": str(rel),
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
            })
    return {"files": files}


@app.get("/api/projects/{project_id}/files/{file_path:path}")
async def get_project_file(project_id: str, file_path: str):
    """Read a file from a project."""
    target = _resolve_project_path(project_id, file_path)
    if target is None:
        return JSONResponse({"error": "Invalid file path"}, status_code=400)
    if not target.is_file():
        return JSONResponse({"error": "File not found"}, status_code=404)

    size = target.stat().st_size
    try:
        content = target.read_text(encoding="utf-8")
        return {"path": file_path, "content": content, "size": size}
    except (UnicodeDecodeError, ValueError):
        return {"path": file_path, "binary": True, "size": size}


@app.post("/api/projects/{project_id}/files/{file_path:path}")
async def write_project_file(project_id: str, file_path: str, request: Request):
    """Write content to a file in a project."""
    target = _resolve_project_path(project_id, file_path)
    if target is None:
        return JSONResponse({"error": "Invalid file path"}, status_code=400)

    project_dir = _resolve_project_path(project_id)
    if project_dir is None or not project_dir.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    body = await request.json()
    content = body.get("content", "")

    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")

    return {"path": file_path, "size": target.stat().st_size}


@app.delete("/api/projects/{project_id}/files/{file_path:path}")
async def delete_project_file(project_id: str, file_path: str):
    """Delete a file from a project."""
    target = _resolve_project_path(project_id, file_path)
    if target is None:
        return JSONResponse({"error": "Invalid file path"}, status_code=400)
    if not target.is_file():
        return JSONResponse({"error": "File not found"}, status_code=404)

    target.unlink()
    return {"ok": True}


@app.get("/api/projects/{project_id}/download/{file_path:path}")
async def download_project_file(project_id: str, file_path: str):
    """Download a file from a project."""
    target = _resolve_project_path(project_id, file_path)
    if target is None:
        return JSONResponse({"error": "Invalid file path"}, status_code=400)
    if not target.is_file():
        return JSONResponse({"error": "File not found"}, status_code=404)

    return FileResponse(
        path=str(target),
        filename=target.name,
        headers={"Content-Disposition": f"attachment; filename=\"{target.name}\""},
    )


@app.get("/api/projects/{project_id}/view/{file_path:path}")
async def view_project_file(project_id: str, file_path: str):
    """Serve a file inline (opens in browser instead of downloading). Useful for PDFs."""
    target = _resolve_project_path(project_id, file_path)
    if target is None:
        return JSONResponse({"error": "Invalid file path"}, status_code=400)
    if not target.is_file():
        return JSONResponse({"error": "File not found"}, status_code=404)

    ext = target.suffix.lower()
    mime_types = {
        ".pdf": "application/pdf",
        ".html": "text/html",
        ".htm": "text/html",
        ".md": "text/markdown",
        ".txt": "text/plain",
        ".json": "application/json",
        ".csv": "text/csv",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".svg": "image/svg+xml",
    }
    media_type = mime_types.get(ext, "application/octet-stream")

    return FileResponse(
        path=str(target),
        filename=target.name,
        media_type=media_type,
        headers={"Content-Disposition": f"inline; filename=\"{target.name}\""},
    )


@app.get("/api/projects/{project_id}/knowledge")
async def get_project_knowledge(project_id: str):
    """Get project knowledge base summary."""
    project_path = PROJECTS_DIR / project_id
    knowledge_file = project_path / "knowledge.json"
    if not knowledge_file.exists():
        return JSONResponse({"error": "Project not found"}, status_code=404)
    try:
        knowledge = json.loads(knowledge_file.read_text(encoding="utf-8"))
        return {
            "name": knowledge.get("name", project_id),
            "findings": knowledge.get("findings", []),
            "sources": knowledge.get("sources", []),
        }
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)


@app.post("/api/projects/{project_id}/research")
async def research_project(project_id: str, request: Request):
    """Autonomous multi-step research agent. Streams SSE progress."""
    body = await request.json()
    topic = (body.get("topic") or "").strip()
    if not topic:
        return JSONResponse({"error": "No topic provided"}, status_code=400)

    depth = body.get("depth", "standard")
    model = body.get("model", FALLBACK_MODEL)
    num_queries = {"quick": 3, "standard": 5, "deep": 8}.get(depth, 5)

    project_path = PROJECTS_DIR / project_id
    if not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    knowledge_file = project_path / "knowledge.json"
    if not knowledge_file.exists():
        return JSONResponse({"error": "Invalid project"}, status_code=404)

    async def research_stream():
        try:
            yield _sse_event({"type": "status", "message": f"Planning research on: {topic}"})
            log.info("Research started: %s depth=%s", topic, depth)

            try:
                queries = await _generate_search_queries(topic, num_queries, model)
            except Exception:
                queries = [topic]
            yield _sse_event({"type": "status", "message": f"Generated {len(queries)} search queries"})

            all_sources = []
            all_page_content = []
            search_errors = 0

            for i, query in enumerate(queries):
                yield _sse_event({"type": "status", "message": f"Searching ({i+1}/{len(queries)}): {query}"})

                try:
                    search_results = await _do_web_search(query, max_results=5)
                except Exception:
                    search_results = []
                    search_errors += 1
                    yield _sse_event({"type": "status", "message": f"⚠ Search failed for: {query}, continuing..."})

                yield _sse_event({"type": "search", "query": query, "results_count": len(search_results)})

                for result in search_results[:2]:
                    url = result.get("href", "")
                    title = result.get("title", "")
                    snippet = result.get("body", "")

                    all_sources.append({
                        "url": url,
                        "title": title,
                        "snippet": snippet,
                        "fetched_at": datetime.now(timezone.utc).isoformat(),
                    })

                    try:
                        page_text = await _fetch_page_content(url)
                        if page_text:
                            all_page_content.append(f"Source: {title} ({url})\n{page_text[:5000]}")
                            yield _sse_event({"type": "source", "message": f"Read: {title}"})
                    except Exception:
                        yield _sse_event({"type": "status", "message": f"⚠ Could not read: {title[:50]}"})

            if not all_sources and not all_page_content:
                yield _sse_event({"type": "error", "error": f"All {len(queries)} searches failed. Check your internet connection."})
                return

            if search_errors > 0:
                yield _sse_event({"type": "status", "message": f"Completed with {search_errors} failed search(es), {len(all_sources)} sources found"})

            yield _sse_event({"type": "status", "message": "Analyzing sources and extracting findings..."})

            try:
                findings_text = await _extract_findings(topic, all_page_content, model)
            except Exception as exc:
                findings_text = "Key findings from search results:\n\n" + "\n\n".join(
                    f"- {s.get('title', 'Unknown')}: {s.get('snippet', '')}" for s in all_sources[:10]
                )
                yield _sse_event({"type": "status", "message": f"⚠ Analysis partially failed ({exc}), using raw findings"})

            yield _sse_event({"type": "finding", "text": findings_text[:2000]})

            yield _sse_event({"type": "status", "message": "Writing comprehensive summary..."})

            # Emit numbered source map so frontend can render clickable citation links
            source_map = [
                {"index": i, "title": s.get("title", "Unknown"), "url": s.get("url", "")}
                for i, s in enumerate(all_sources[:15], 1)
            ]
            yield _sse_event({"type": "source_map", "sources": source_map})

            try:
                summary = await _synthesize_summary(topic, findings_text, all_sources, model)
            except Exception as exc:
                summary = f"# Research: {topic}\n\n## Findings\n\n{findings_text}\n\n## Sources\n\n" + "\n".join(
                    f"- [{s.get('title', 'Unknown')}]({s.get('url', '')})" for s in all_sources[:10]
                )
                yield _sse_event({"type": "status", "message": f"⚠ Summary generation failed ({exc}), using raw findings"})

            yield _sse_event({"type": "status", "message": "Saving research to project..."})

            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            topic_slug = _slugify(topic)[:40]
            note_filename = f"{timestamp}-{topic_slug}.md"
            notes_dir = project_path / "notes"
            notes_dir.mkdir(exist_ok=True)
            note_path = notes_dir / note_filename
            note_path.write_text(summary, encoding="utf-8")

            pdf_filename = f"{timestamp}-{topic_slug}.pdf"
            try:
                await asyncio.to_thread(_save_markdown_as_pdf, summary, notes_dir / pdf_filename, topic)
            except Exception:
                pass

            knowledge = json.loads(knowledge_file.read_text(encoding="utf-8"))

            existing_urls = {s["url"] for s in knowledge.get("sources", [])}
            for src in all_sources:
                if src["url"] not in existing_urls:
                    knowledge.setdefault("sources", []).append(src)
                    existing_urls.add(src["url"])

            knowledge.setdefault("findings", []).append({
                "topic": topic,
                "summary": findings_text[:2000],
                "sources": [s["url"] for s in all_sources[:10]],
                "timestamp": datetime.now(timezone.utc).isoformat(),
            })

            knowledge_file.write_text(json.dumps(knowledge, indent=2, ensure_ascii=False), encoding="utf-8")

            log.info("Research complete: %s — %d sources", topic, len(all_sources))
            yield _sse_event({"type": "content", "text": summary})

            yield _sse_event({
                "type": "done",
                "message": "Research complete!",
                "summary_file": f"notes/{note_filename}",
            })

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        research_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


def _parse_code_files(content: str) -> list[tuple[str, str]]:
    """Parse multi-file code output using === FILE: path === / === END FILE === markers."""
    files = []
    pattern = re.compile(
        r'===\s*FILE:\s*(.+?)\s*===\n(.*?)\n===\s*END\s*FILE\s*===',
        re.DOTALL,
    )
    for match in pattern.finditer(content):
        file_path = match.group(1).strip()
        file_content = match.group(2)
        if file_path and file_content is not None:
            files.append((file_path, file_content))
    return files


def _parse_workflow_plan(text: str) -> list[dict]:
    """Parse a JSON workflow plan from LLM output."""
    text = text.strip()
    if "```" in text:
        lines = text.split("\n")
        in_fence = False
        json_lines = []
        for line in lines:
            if line.strip().startswith("```"):
                in_fence = not in_fence
                continue
            if in_fence:
                json_lines.append(line)
        if json_lines:
            text = "\n".join(json_lines)

    start = text.find("[")
    end = text.rfind("]")
    if start >= 0 and end > start:
        text = text[start:end + 1]

    try:
        steps = json.loads(text)
        if isinstance(steps, list):
            valid_types = {"research", "document", "code", "data"}
            return [s for s in steps if isinstance(s, dict) and s.get("type") in valid_types]
    except (json.JSONDecodeError, ValueError):
        pass
    return []


@app.post("/api/projects/{project_id}/generate-document")
async def generate_document(project_id: str, request: Request):
    """Generate a Markdown document, optionally using project knowledge. Streams SSE."""
    body = await request.json()
    topic = (body.get("topic") or "").strip()
    if not topic:
        return JSONResponse({"error": "No topic provided"}, status_code=400)

    project_path = PROJECTS_DIR / project_id
    if not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    doc_type = body.get("type", "report")
    model = body.get("model", FALLBACK_MODEL)
    use_knowledge = body.get("use_knowledge", True)

    async def doc_stream():
        try:
            yield _sse_event({"type": "status", "message": f"Generating {doc_type}: {topic}"})

            context = ""
            if use_knowledge:
                context = _get_project_knowledge_context(project_id, max_chars=6000)

            prompt = f"""Write a comprehensive {doc_type} about: "{topic}"

{f'Use this research context to inform your writing:{chr(10)}{context}' if context else ''}

Format as a well-structured Markdown document with:
- A clear title (# heading)
- An introduction/executive summary
- Organized sections with ## subheadings
- Bullet points and numbered lists where appropriate
- A conclusion or summary section
- A "## References" section at the end listing all cited sources as a numbered list

If sources are provided above, use inline citation markers [1], [2], etc. to reference them. Place the citation immediately after the claim it supports. In the References section, list each cited source as: [N] Title — URL. Only cite sources from the provided list.

Be thorough, detailed, and well-organized."""

            yield _sse_event({"type": "status", "message": "Writing document..."})

            async with httpx.AsyncClient(timeout=120.0) as client:
                resp = await client.post(f"{OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                    "options": {"temperature": 0.5, "num_predict": 8192, "num_ctx": 32768},
                })
                data = resp.json()
                content = data.get("message", {}).get("content", "")

            if not content:
                yield _sse_event({"type": "error", "error": "Model returned empty content"})
                return

            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            topic_slug = _slugify(topic)[:40]
            filename = f"{timestamp}-{topic_slug}.md"
            output_dir = project_path / "output"
            output_dir.mkdir(exist_ok=True)
            file_path = output_dir / filename
            file_path.write_text(content, encoding="utf-8")
            log.info("Document generated: %s", filename)

            pdf_filename = f"{timestamp}-{topic_slug}.pdf"
            pdf_path = output_dir / pdf_filename
            try:
                await asyncio.to_thread(_save_markdown_as_pdf, content, pdf_path, topic)
                yield _sse_event({"type": "status", "message": f"PDF saved: output/{pdf_filename}"})
            except Exception:
                pass

            yield _sse_event({"type": "content", "text": content})
            yield _sse_event({"type": "done", "message": "Document generated", "file_path": f"output/{filename}"})

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        doc_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/api/projects/{project_id}/export-pdf")
async def export_pdf(project_id: str, request: Request):
    """Convert a Markdown file to a human-readable PDF."""
    body = await request.json()
    file_path = (body.get("file_path") or "").strip()
    if not file_path:
        return JSONResponse({"error": "No file_path provided"}, status_code=400)

    resolved = _resolve_project_path(project_id, file_path)
    if resolved is None or not resolved.exists():
        return JSONResponse({"error": "File not found"}, status_code=404)

    md_content = resolved.read_text(encoding="utf-8")

    if _MARKDOWN_AVAILABLE:
        html_body = _markdown_lib.markdown(md_content, extensions=["tables", "fenced_code", "toc", "nl2br"])
    else:
        html_body = f"<pre>{md_content}</pre>"

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{Path(file_path).stem}</title>
<style>
@page {{ size: A4; margin: 2.5cm; }}
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Helvetica Neue', sans-serif; font-size: 11pt; line-height: 1.7; color: #1a1a1a; }}
h1 {{ font-size: 22pt; border-bottom: 2px solid #333; padding-bottom: 8px; margin-top: 0; margin-bottom: 16px; }}
h2 {{ font-size: 16pt; border-bottom: 1px solid #ccc; padding-bottom: 6px; margin-top: 28px; color: #2c3e50; }}
h3 {{ font-size: 13pt; margin-top: 20px; color: #34495e; }}
p {{ margin: 8px 0; text-align: justify; }}
ul, ol {{ margin: 8px 0; padding-left: 24px; }}
li {{ margin: 4px 0; }}
code {{ background: #f4f4f4; padding: 2px 5px; border-radius: 3px; font-size: 0.88em; font-family: 'SF Mono', 'Fira Code', Menlo, monospace; }}
pre {{ background: #f8f8f8; padding: 14px 18px; border-radius: 6px; overflow-x: auto; border: 1px solid #e8e8e8; font-size: 0.85em; line-height: 1.5; }}
pre code {{ background: none; padding: 0; }}
table {{ border-collapse: collapse; width: 100%; margin: 14px 0; font-size: 0.92em; }}
th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
th {{ background: #f0f0f0; font-weight: 600; }}
tr:nth-child(even) {{ background: #fafafa; }}
blockquote {{ border-left: 4px solid #3498db; margin: 14px 0; padding: 10px 18px; color: #555; background: #f9fbfd; border-radius: 0 4px 4px 0; }}
a {{ color: #2980b9; text-decoration: none; }}
hr {{ border: none; border-top: 1px solid #ddd; margin: 20px 0; }}
</style></head><body>{html_body}</body></html>"""

    if _WEASYPRINT_AVAILABLE:
        def _generate_pdf():
            doc = _weasyprint.HTML(string=html)
            return doc.write_pdf()

        pdf_bytes = await asyncio.to_thread(_generate_pdf)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{Path(file_path).stem}.pdf"'},
        )

    return Response(
        content=html.encode("utf-8"),
        media_type="text/html",
        headers={"Content-Disposition": f'attachment; filename="{Path(file_path).stem}.html"'},
    )


@app.post("/api/projects/{project_id}/export-docx")
async def export_docx(project_id: str, request: Request):
    """Convert a Markdown file to a DOCX document."""
    body = await request.json()
    file_path = (body.get("file_path") or "").strip()
    if not file_path:
        return JSONResponse({"error": "No file_path provided"}, status_code=400)

    resolved = _resolve_project_path(project_id, file_path)
    if resolved is None or not resolved.exists():
        return JSONResponse({"error": "File not found"}, status_code=404)

    if not _DOCX_AVAILABLE:
        return JSONResponse({"error": "python-docx is not installed"}, status_code=500)

    md_content = resolved.read_text(encoding="utf-8")

    def _build_docx(md_text: str) -> bytes:
        doc = _docx_module.Document()

        in_code_block = False
        code_lines: list[str] = []

        for line in md_text.split("\n"):
            # --- fenced code block toggle ---
            if line.strip().startswith("```"):
                if in_code_block:
                    # close code block
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = _docx_module.shared.Pt(9)
                    p.style = doc.styles["Normal"]
                    fmt = p.paragraph_format
                    fmt.space_before = _docx_module.shared.Pt(4)
                    fmt.space_after = _docx_module.shared.Pt(4)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            stripped = line.strip()

            # --- blank lines ---
            if not stripped:
                continue

            # --- headings ---
            heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
            if heading_match:
                level = min(len(heading_match.group(1)), 3)
                doc.add_heading(heading_match.group(2).strip(), level=level)
                continue

            # --- bullet lists ---
            bullet_match = re.match(r"^[\-\*]\s+(.*)", stripped)
            if bullet_match:
                doc.add_paragraph(bullet_match.group(1), style="List Bullet")
                continue

            # --- numbered lists ---
            num_match = re.match(r"^\d+[\.\)]\s+(.*)", stripped)
            if num_match:
                doc.add_paragraph(num_match.group(1), style="List Number")
                continue

            # --- horizontal rule ---
            if re.match(r"^[-*_]{3,}\s*$", stripped):
                doc.add_paragraph("─" * 40)
                continue

            # --- blockquote ---
            if stripped.startswith(">"):
                text = re.sub(r"^>\s?", "", stripped)
                p = doc.add_paragraph(text)
                fmt = p.paragraph_format
                fmt.left_indent = _docx_module.shared.Inches(0.5)
                continue

            # --- normal paragraph ---
            doc.add_paragraph(stripped)

        # flush any unclosed code block
        if code_lines:
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = _docx_module.shared.Pt(9)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    docx_bytes = await asyncio.to_thread(_build_docx, md_content)
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{Path(file_path).stem}.docx"'},
    )


@app.post("/api/projects/{project_id}/export-html")
async def export_html(project_id: str, request: Request):
    """Convert a Markdown file to a standalone styled HTML document."""
    body = await request.json()
    file_path = (body.get("file_path") or "").strip()
    if not file_path:
        return JSONResponse({"error": "No file_path provided"}, status_code=400)

    resolved = _resolve_project_path(project_id, file_path)
    if resolved is None or not resolved.exists():
        return JSONResponse({"error": "File not found"}, status_code=404)

    md_content = resolved.read_text(encoding="utf-8")

    if _MARKDOWN_AVAILABLE:
        html_body = _markdown_lib.markdown(md_content, extensions=["tables", "fenced_code", "toc", "nl2br"])
    else:
        html_body = f"<pre>{md_content}</pre>"

    html = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>{Path(file_path).stem}</title>
<style>
body {{ max-width: 860px; margin: 40px auto; padding: 0 24px; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Helvetica Neue', sans-serif; font-size: 16px; line-height: 1.7; color: #1a1a1a; background: #fff; }}
h1 {{ font-size: 2em; border-bottom: 2px solid #333; padding-bottom: 8px; margin-top: 0; margin-bottom: 16px; }}
h2 {{ font-size: 1.5em; border-bottom: 1px solid #ccc; padding-bottom: 6px; margin-top: 32px; color: #2c3e50; }}
h3 {{ font-size: 1.17em; margin-top: 24px; color: #34495e; }}
p {{ margin: 10px 0; }}
ul, ol {{ margin: 10px 0; padding-left: 28px; }}
li {{ margin: 4px 0; }}
code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; font-family: 'SF Mono', 'Fira Code', Menlo, monospace; }}
pre {{ background: #f8f8f8; padding: 16px 20px; border-radius: 6px; overflow-x: auto; border: 1px solid #e8e8e8; font-size: 0.88em; line-height: 1.5; }}
pre code {{ background: none; padding: 0; }}
table {{ border-collapse: collapse; width: 100%; margin: 16px 0; font-size: 0.95em; }}
th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
th {{ background: #f0f0f0; font-weight: 600; }}
tr:nth-child(even) {{ background: #fafafa; }}
blockquote {{ border-left: 4px solid #3498db; margin: 16px 0; padding: 10px 20px; color: #555; background: #f9fbfd; border-radius: 0 4px 4px 0; }}
a {{ color: #2980b9; text-decoration: none; }}
a:hover {{ text-decoration: underline; }}
hr {{ border: none; border-top: 1px solid #ddd; margin: 24px 0; }}
img {{ max-width: 100%; height: auto; }}
@media (prefers-color-scheme: dark) {{
  body {{ background: #1a1a2e; color: #e0e0e0; }}
  h1 {{ border-bottom-color: #555; }}
  h2 {{ border-bottom-color: #444; color: #8ab4f8; }}
  h3 {{ color: #8ab4f8; }}
  code {{ background: #2a2a3e; }}
  pre {{ background: #252540; border-color: #333; }}
  th {{ background: #2a2a3e; }}
  tr:nth-child(even) {{ background: #1f1f35; }}
  td, th {{ border-color: #333; }}
  blockquote {{ background: #1f1f35; border-left-color: #3498db; color: #aaa; }}
  a {{ color: #5dade2; }}
}}
</style></head><body>{html_body}</body></html>"""

    return Response(
        content=html.encode("utf-8"),
        media_type="text/html",
        headers={"Content-Disposition": f'attachment; filename="{Path(file_path).stem}.html"'},
    )


@app.post("/api/projects/{project_id}/generate-data")
async def generate_data(project_id: str, request: Request):
    """Generate structured data (CSV or JSON). Streams SSE."""
    body = await request.json()
    topic = (body.get("topic") or "").strip()
    if not topic:
        return JSONResponse({"error": "No topic provided"}, status_code=400)

    project_path = PROJECTS_DIR / project_id
    if not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    data_format = body.get("format", "csv")
    model = body.get("model", FALLBACK_MODEL)

    async def data_stream():
        try:
            yield _sse_event({"type": "status", "message": f"Generating {data_format.upper()} data: {topic}"})

            if data_format == "csv":
                prompt = f"""Generate a CSV dataset about: "{topic}"

Rules:
- First row must be column headers
- Use commas as delimiters
- Wrap fields containing commas in double quotes
- Include at least 10 rows of meaningful data
- Return ONLY the CSV content, no explanations or markdown formatting"""
            else:
                prompt = f"""Generate a JSON dataset about: "{topic}"

Rules:
- Return a JSON array of objects
- Each object should have consistent keys
- Include at least 10 items with meaningful data
- Return ONLY valid JSON, no explanations or markdown formatting"""

            async with httpx.AsyncClient(timeout=60.0) as client:
                resp = await client.post(f"{OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                    "options": {"temperature": 0.3, "num_predict": 4096, "num_ctx": 32768},
                })
                data = resp.json()
                content = data.get("message", {}).get("content", "")

            if not content:
                yield _sse_event({"type": "error", "error": "Model returned empty content"})
                return

            content = content.strip()
            if content.startswith("```"):
                lines = content.split("\n")
                if lines[-1].strip() == "```":
                    lines = lines[1:-1]
                else:
                    lines = lines[1:]
                content = "\n".join(lines)

            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            topic_slug = _slugify(topic)[:30]
            ext = "csv" if data_format == "csv" else "json"
            filename = f"{timestamp}-{topic_slug}.{ext}"
            data_dir = project_path / "output" / "data"
            data_dir.mkdir(parents=True, exist_ok=True)
            file_path = data_dir / filename
            file_path.write_text(content, encoding="utf-8")

            yield _sse_event({"type": "content", "text": content, "format": data_format})
            yield _sse_event({"type": "done", "message": f"{data_format.upper()} generated", "file_path": f"output/data/{filename}"})

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        data_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/api/projects/{project_id}/generate-code")
async def generate_code(project_id: str, request: Request):
    """Generate a multi-file code project. Streams SSE."""
    body = await request.json()
    description = (body.get("description") or "").strip()
    if not description:
        return JSONResponse({"error": "No description provided"}, status_code=400)

    project_path = PROJECTS_DIR / project_id
    if not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    model = body.get("model", FALLBACK_MODEL)

    async def code_stream():
        try:
            yield _sse_event({"type": "status", "message": f"Generating code: {description[:80]}"})

            prompt = f"""Generate a complete code project for: "{description}"

IMPORTANT: For EACH file, use this EXACT format:
=== FILE: path/to/filename.ext ===
<file content here>
=== END FILE ===

Rules:
- Generate all necessary files for a working project
- Include a README.md with setup instructions
- Use best practices and proper project structure
- Include basic error handling
- Add comments where helpful
- Make sure the code is complete and runnable"""

            yield _sse_event({"type": "status", "message": "Writing code..."})

            async with httpx.AsyncClient(timeout=120.0) as client:
                resp = await client.post(f"{OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                    "options": {"temperature": 0.3, "num_predict": 8192, "num_ctx": 32768},
                })
                data = resp.json()
                content = data.get("message", {}).get("content", "")

            if not content:
                yield _sse_event({"type": "error", "error": "Model returned empty content"})
                return

            files = _parse_code_files(content)

            if not files:
                timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
                slug = _slugify(description)[:30]
                filename = f"{timestamp}-{slug}.txt"
                code_dir = project_path / "output" / "code"
                code_dir.mkdir(parents=True, exist_ok=True)
                (code_dir / filename).write_text(content, encoding="utf-8")
                yield _sse_event({"type": "file", "path": f"output/code/{filename}", "size": len(content)})
                yield _sse_event({"type": "done", "message": "Code generated (single file)", "files_count": 1})
                return

            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            slug = _slugify(description)[:30]
            code_dir = project_path / "output" / "code" / f"{timestamp}-{slug}"
            code_dir.mkdir(parents=True, exist_ok=True)

            saved_files = []
            for file_path, file_content in files:
                safe_path = Path(file_path.lstrip("/").replace("..", ""))
                full_path = code_dir / safe_path
                full_path.parent.mkdir(parents=True, exist_ok=True)
                full_path.write_text(file_content, encoding="utf-8")
                rel_path = str(full_path.relative_to(project_path))
                saved_files.append(rel_path)
                yield _sse_event({"type": "file", "path": rel_path, "size": len(file_content)})

            yield _sse_event({"type": "done", "message": f"Generated {len(saved_files)} files", "files_count": len(saved_files)})

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        code_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/api/projects/{project_id}/workflow")
async def run_workflow(project_id: str, request: Request):
    """Execute a multi-step workflow. Streams SSE progress."""
    body = await request.json()
    user_request = (body.get("request") or "").strip()
    if not user_request:
        return JSONResponse({"error": "No request provided"}, status_code=400)

    project_path = PROJECTS_DIR / project_id
    if not project_path.is_dir():
        return JSONResponse({"error": "Project not found"}, status_code=404)

    model = body.get("model", FALLBACK_MODEL)

    async def workflow_stream():
        try:
            yield _sse_event({"type": "status", "message": "Planning workflow..."})

            plan_prompt = f"""Analyze this request and break it into ordered steps. For each step, specify the type.

Request: "{user_request}"

Available step types:
- research: Search the web and gather information on a topic
- document: Write a document/report about a topic
- code: Generate code for a programming task
- data: Generate structured data (CSV/JSON tables)

Return ONLY a JSON array of steps, no explanation. Example:
[{{"type": "research", "description": "Research topic X"}}, {{"type": "document", "description": "Write a report on X"}}]"""

            async with httpx.AsyncClient(timeout=30.0) as client:
                resp = await client.post(f"{OLLAMA}/api/chat", json={
                    "model": model,
                    "messages": [{"role": "user", "content": plan_prompt}],
                    "stream": False,
                    "options": {"temperature": 0.2, "num_predict": 1024},
                })
                plan_data = resp.json()
                plan_text = plan_data.get("message", {}).get("content", "")

            steps = _parse_workflow_plan(plan_text)
            if not steps:
                steps = [{"type": "document", "description": user_request}]

            yield _sse_event({"type": "plan", "steps": steps})

            for i, step in enumerate(steps):
                step_type = step.get("type", "document")
                step_desc = step.get("description", user_request)

                yield _sse_event({
                    "type": "step_start",
                    "step": i + 1,
                    "total": len(steps),
                    "step_type": step_type,
                    "description": step_desc,
                })

                if step_type == "research":
                    queries = await _generate_search_queries(step_desc, 3, model)
                    all_sources = []
                    all_content = []
                    for query in queries:
                        results = await _do_web_search(query, max_results=3)
                        for r in results[:2]:
                            all_sources.append(r)
                            page = await _fetch_page_content(r.get("href", ""))
                            if page:
                                all_content.append(page[:2000])

                    if all_content:
                        findings = await _extract_findings(step_desc, all_content, model)
                        knowledge_file = project_path / "knowledge.json"
                        knowledge = json.loads(knowledge_file.read_text(encoding="utf-8"))
                        knowledge.setdefault("findings", []).append({
                            "topic": step_desc,
                            "summary": findings[:2000],
                            "sources": [s.get("href", "") for s in all_sources[:10]],
                            "timestamp": datetime.now(timezone.utc).isoformat(),
                        })
                        knowledge_file.write_text(json.dumps(knowledge, indent=2, ensure_ascii=False), encoding="utf-8")

                    yield _sse_event({"type": "step_done", "step": i + 1, "message": f"Research complete: {len(all_sources)} sources found"})

                elif step_type == "document":
                    context = _get_project_knowledge_context(project_id, max_chars=6000)
                    doc_prompt = f"Write a comprehensive document about: {step_desc}\n\n{context if context else ''}"

                    async with httpx.AsyncClient(timeout=120.0) as client:
                        resp = await client.post(f"{OLLAMA}/api/chat", json={
                            "model": model,
                            "messages": [{"role": "user", "content": doc_prompt}],
                            "stream": False,
                            "options": {"temperature": 0.5, "num_predict": 8192, "num_ctx": 32768},
                        })
                        doc_data = resp.json()
                        doc_content = doc_data.get("message", {}).get("content", "")

                    if doc_content:
                        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
                        slug = _slugify(step_desc)[:30]
                        output_dir = project_path / "output"
                        output_dir.mkdir(exist_ok=True)
                        (output_dir / f"{timestamp}-{slug}.md").write_text(doc_content, encoding="utf-8")

                    yield _sse_event({"type": "step_done", "step": i + 1, "message": "Document generated"})

                elif step_type == "code":
                    code_prompt = f"""Generate a complete code project for: "{step_desc}"

For EACH file use: === FILE: path === ... === END FILE ==="""

                    async with httpx.AsyncClient(timeout=120.0) as client:
                        resp = await client.post(f"{OLLAMA}/api/chat", json={
                            "model": model,
                            "messages": [{"role": "user", "content": code_prompt}],
                            "stream": False,
                            "options": {"temperature": 0.3, "num_predict": 8192, "num_ctx": 32768},
                        })
                        code_data = resp.json()
                        code_content = code_data.get("message", {}).get("content", "")

                    if code_content:
                        files = _parse_code_files(code_content)
                        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
                        slug = _slugify(step_desc)[:30]
                        code_dir = project_path / "output" / "code" / f"{timestamp}-{slug}"
                        code_dir.mkdir(parents=True, exist_ok=True)
                        if files:
                            for fp, fc in files:
                                safe = Path(fp.lstrip("/").replace("..", ""))
                                full = code_dir / safe
                                full.parent.mkdir(parents=True, exist_ok=True)
                                full.write_text(fc, encoding="utf-8")
                        else:
                            (code_dir / "output.txt").write_text(code_content, encoding="utf-8")

                    yield _sse_event({"type": "step_done", "step": i + 1, "message": "Code generated"})

                elif step_type == "data":
                    data_prompt = f"""Generate a CSV dataset about: "{step_desc}"
Return ONLY CSV content with headers. No explanations."""

                    async with httpx.AsyncClient(timeout=60.0) as client:
                        resp = await client.post(f"{OLLAMA}/api/chat", json={
                            "model": model,
                            "messages": [{"role": "user", "content": data_prompt}],
                            "stream": False,
                            "options": {"temperature": 0.3, "num_predict": 4096, "num_ctx": 32768},
                        })
                        csv_data = resp.json()
                        csv_content = csv_data.get("message", {}).get("content", "")

                    if csv_content:
                        csv_content = csv_content.strip()
                        if csv_content.startswith("```"):
                            lines = csv_content.split("\n")
                            if lines[-1].strip() == "```":
                                lines = lines[1:-1]
                            else:
                                lines = lines[1:]
                            csv_content = "\n".join(lines)

                        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
                        slug = _slugify(step_desc)[:30]
                        data_dir = project_path / "output" / "data"
                        data_dir.mkdir(parents=True, exist_ok=True)
                        (data_dir / f"{timestamp}-{slug}.csv").write_text(csv_content, encoding="utf-8")

                    yield _sse_event({"type": "step_done", "step": i + 1, "message": "Data generated"})

                else:
                    yield _sse_event({"type": "step_done", "step": i + 1, "message": f"Skipped unknown step type: {step_type}"})

            yield _sse_event({"type": "done", "message": f"Workflow complete! Executed {len(steps)} steps."})

        except Exception as exc:
            yield _sse_event({"type": "error", "error": str(exc)})

    return StreamingResponse(
        workflow_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.get("/api/tools")
async def list_tools():
    return {"tools": _load_tool_registry()}

@app.get("/api/tools/{tool_name}")
async def get_tool(tool_name: str):
    registry = _load_tool_registry()
    tool = next((t for t in registry if t["name"] == tool_name), None)
    if not tool:
        return JSONResponse({"error": "Tool not found"}, status_code=404)
    module_path = PLUGINS_DIR / tool["module"]
    code = module_path.read_text(encoding="utf-8") if module_path.exists() else ""
    return {**tool, "code": code}

@app.post("/api/tools/{tool_name}/execute")
async def execute_tool_endpoint(tool_name: str, request: Request):
    body = await request.json()
    params = body.get("params", {})
    result = await asyncio.to_thread(_execute_tool, tool_name, params)
    return result

@app.delete("/api/tools/{tool_name}")
async def delete_tool(tool_name: str):
    registry = _load_tool_registry()
    tool = next((t for t in registry if t["name"] == tool_name), None)
    if not tool:
        return JSONResponse({"error": "Tool not found"}, status_code=404)
    module_path = PLUGINS_DIR / tool["module"]
    if module_path.exists():
        module_path.unlink()
    registry = [t for t in registry if t["name"] != tool_name]
    _save_tool_registry(registry)
    return {"ok": True}

@app.post("/api/tools/{tool_name}/toggle")
async def toggle_tool(tool_name: str):
    registry = _load_tool_registry()
    for t in registry:
        if t["name"] == tool_name:
            t["enabled"] = not t.get("enabled", True)
            _save_tool_registry(registry)
            return {"ok": True, "enabled": t["enabled"]}
    return JSONResponse({"error": "Tool not found"}, status_code=404)

@app.post("/api/tools/{tool_name}/preview")
async def preview_tool(tool_name: str):
    """Return tool source, parameters, and sample invocation without executing."""
    registry = _load_tool_registry()
    tool = next((t for t in registry if t["name"] == tool_name), None)
    if not tool:
        return JSONResponse({"error": "Tool not found"}, status_code=404)
    module_path = PLUGINS_DIR / tool["module"]
    code = module_path.read_text(encoding="utf-8") if module_path.exists() else ""
    sample_params = {}
    for k, v in tool.get("parameters", {}).items():
        if v.get("type") == "number":
            sample_params[k] = 0
        elif v.get("type") == "boolean":
            sample_params[k] = True
        else:
            sample_params[k] = f"<{v.get('description', k)}>"
    return {
        "name": tool["name"],
        "description": tool.get("description", ""),
        "parameters": tool.get("parameters", {}),
        "code": code,
        "enabled": tool.get("enabled", True),
        "sample_invocation": {
            "endpoint": f"/api/tools/{tool_name}/execute",
            "method": "POST",
            "body": {"params": sample_params},
        },
    }

@app.post("/api/tools/build")
async def build_tool_endpoint(request: Request):
    body = await request.json()
    description = (body.get("description") or "").strip()
    model = body.get("model", FALLBACK_MODEL)
    preview = body.get("preview", False)
    if not description:
        return JSONResponse({"error": "No description provided"}, status_code=400)
    if preview:
        result = await _build_tool_preview(description, model)
    else:
        result = await _build_tool(description, model)
    return result

@app.get("/api/memory")
async def get_memories():
    return {"memories": _load_memories()}

@app.post("/api/memory")
async def add_memory_endpoint(request: Request):
    body = await request.json()
    text = (body.get("text") or "").strip()
    if not text:
        return JSONResponse({"error": "No memory text provided"}, status_code=400)
    _add_memory(text)
    return {"ok": True, "memories": _load_memories()}

@app.delete("/api/memory/{index}")
async def remove_memory_endpoint(index: int):
    if _remove_memory(index):
        return {"ok": True, "memories": _load_memories()}
    return JSONResponse({"error": "Invalid index"}, status_code=404)

@app.post("/api/suggest-followups")
async def suggest_followups(request: Request):
    """Generate follow-up question suggestions based on the conversation."""
    body = await request.json()
    model = body.get("model", FALLBACK_MODEL)
    messages_list = body.get("messages", [])
    if len(messages_list) < 2:
        return {"suggestions": []}

    last_exchange = messages_list[-2:]
    context = "\n".join(f"{m['role']}: {m['content'][:300]}" for m in last_exchange)

    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            resp = await client.post(f"{OLLAMA}/api/chat", json={
                "model": model,
                "messages": [{
                    "role": "user",
                    "content": f"Based on this conversation, suggest exactly 3 short follow-up questions the user might ask next. Return ONLY the questions, one per line, no numbering.\n\n{context}",
                }],
                "stream": False,
                "options": {"temperature": 0.7, "num_predict": 128},
            })
            data = resp.json()
            content = data.get("message", {}).get("content", "")
            suggestions = [q.strip().strip('"').strip("'") for q in content.strip().split("\n") if q.strip()]
            suggestions = [re.sub(r'^[\d]+[.)\s]+|^[-*]\s+', '', s).strip() for s in suggestions]
            return {"suggestions": suggestions[:3]}
    except Exception:
        return {"suggestions": []}


# ── Data Portability (Export / Import Archive) ──────────────────────────

import zipfile as _zipfile


def _add_dir_to_zip(zf: _zipfile.ZipFile, directory: Path, archive_prefix: str) -> int:
    """Recursively add a directory to a zip file. Returns number of files added."""
    count = 0
    if not directory.is_dir():
        return count
    for fp in sorted(directory.rglob("*")):
        if fp.is_file():
            arcname = f"{archive_prefix}/{fp.relative_to(directory)}"
            zf.write(fp, arcname)
            count += 1
    return count


@app.get("/api/export-archive")
async def export_archive(request: Request):
    """Create a ZIP archive of all user data for portability."""
    if not _runtime_control_allowed(request):
        return JSONResponse(
            {"error": "Export requires localhost access or a valid LAN token."},
            status_code=403,
        )

    def _build_archive() -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".zip", delete=False)
        tmp.close()
        with _zipfile.ZipFile(tmp.name, "w", _zipfile.ZIP_DEFLATED) as zf:
            # Projects
            projects_count = _add_dir_to_zip(zf, PROJECTS_DIR, "projects")

            # Plugins
            plugins_count = _add_dir_to_zip(zf, PLUGINS_DIR, "plugins")

            # Memory
            memory_count = _add_dir_to_zip(zf, MEMORY_DIR, "memory")

            # Token stats
            token_stats_count = 0
            if _TOKEN_STATS_FILE.exists():
                zf.write(_TOKEN_STATS_FILE, "token_stats.json")
                token_stats_count = 1

            # Manifest
            manifest = {
                "version": "1.0",
                "app": "OfflineAI",
                "exported_at": datetime.now(timezone.utc).isoformat(),
                "platform": platform.system(),
                "contents": {
                    "projects_files": projects_count,
                    "plugins_files": plugins_count,
                    "memory_files": memory_count,
                    "token_stats": token_stats_count > 0,
                },
            }
            zf.writestr("manifest.json", json.dumps(manifest, indent=2))
        return tmp.name

    zip_path = await asyncio.to_thread(_build_archive)
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    filename = f"OfflineAI-backup-{timestamp}.zip"

    async def _stream_and_cleanup():
        try:
            with open(zip_path, "rb") as f:
                while True:
                    chunk = f.read(65536)
                    if not chunk:
                        break
                    yield chunk
        finally:
            Path(zip_path).unlink(missing_ok=True)

    return StreamingResponse(
        _stream_and_cleanup(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/import-archive")
async def import_archive(request: Request, file: UploadFile = File(...)):
    """Import a ZIP archive, merging data without overwriting existing entries."""
    if not _runtime_control_allowed(request):
        return JSONResponse(
            {"error": "Import requires localhost access or a valid LAN token."},
            status_code=403,
        )

    content = await file.read()
    tmp = tempfile.NamedTemporaryFile(suffix=".zip", delete=False)
    tmp.write(content)
    tmp.close()

    def _do_import() -> dict:
        summary: dict = {
            "projects_imported": 0,
            "projects_skipped": 0,
            "plugins_imported": 0,
            "memory_merged": 0,
            "token_stats_merged": False,
            "errors": [],
        }
        try:
            with _zipfile.ZipFile(tmp.name, "r") as zf:
                # Validate manifest
                if "manifest.json" not in zf.namelist():
                    summary["errors"].append("Missing manifest.json — not a valid OfflineAI archive.")
                    return summary
                try:
                    manifest = json.loads(zf.read("manifest.json"))
                    if manifest.get("app") != "OfflineAI":
                        summary["errors"].append("Archive is not from OfflineAI.")
                        return summary
                except (json.JSONDecodeError, KeyError) as exc:
                    summary["errors"].append(f"Invalid manifest: {exc}")
                    return summary

                names = zf.namelist()

                # Import projects (merge, don't overwrite)
                for name in names:
                    if name.startswith("projects/") and not name.endswith("/"):
                        rel = name[len("projects/"):]
                        target = PROJECTS_DIR / rel
                        # Security: prevent path traversal
                        try:
                            target.resolve().relative_to(PROJECTS_DIR.resolve())
                        except ValueError:
                            continue
                        if target.exists():
                            # Check if this is a new project directory
                            parts = Path(rel).parts
                            if len(parts) >= 1:
                                project_dir = PROJECTS_DIR / parts[0]
                                if project_dir.exists():
                                    summary["projects_skipped"] += 1
                                    continue
                        target.parent.mkdir(parents=True, exist_ok=True)
                        target.write_bytes(zf.read(name))
                        summary["projects_imported"] += 1

                # Import plugins (merge registry)
                for name in names:
                    if name.startswith("plugins/") and not name.endswith("/"):
                        rel = name[len("plugins/"):]
                        target = PLUGINS_DIR / rel
                        try:
                            target.resolve().relative_to(PLUGINS_DIR.resolve())
                        except ValueError:
                            continue
                        if rel == "registry.json":
                            # Merge registries
                            try:
                                incoming_reg = json.loads(zf.read(name))
                                existing_reg = _load_tool_registry()
                                existing_names = {t["name"] for t in existing_reg}
                                added = 0
                                for tool in incoming_reg:
                                    if tool.get("name") and tool["name"] not in existing_names:
                                        existing_reg.append(tool)
                                        added += 1
                                _save_tool_registry(existing_reg)
                                summary["plugins_imported"] += added
                            except Exception as exc:
                                summary["errors"].append(f"Plugin registry merge error: {exc}")
                        else:
                            # Copy tool module files if not existing
                            if not target.exists():
                                target.parent.mkdir(parents=True, exist_ok=True)
                                target.write_bytes(zf.read(name))

                # Import memory (merge)
                for name in names:
                    if name.startswith("memory/") and not name.endswith("/"):
                        rel = name[len("memory/"):]
                        if rel == "preferences.json":
                            try:
                                incoming_mems = json.loads(zf.read(name))
                                existing_mems = _load_memories()
                                existing_set = set(existing_mems)
                                added = 0
                                for m in incoming_mems:
                                    if isinstance(m, str) and m.strip() and m not in existing_set:
                                        existing_mems.append(m)
                                        existing_set.add(m)
                                        added += 1
                                _save_memories(existing_mems)
                                summary["memory_merged"] = added
                            except Exception as exc:
                                summary["errors"].append(f"Memory merge error: {exc}")
                        else:
                            target = MEMORY_DIR / rel
                            try:
                                target.resolve().relative_to(MEMORY_DIR.resolve())
                            except ValueError:
                                continue
                            if not target.exists():
                                target.parent.mkdir(parents=True, exist_ok=True)
                                target.write_bytes(zf.read(name))

                # Import token stats (merge)
                if "token_stats.json" in names:
                    try:
                        incoming_stats = json.loads(zf.read("token_stats.json"))
                        global _token_stats
                        for model_key, counts in incoming_stats.items():
                            if isinstance(counts, list) and len(counts) == 2:
                                if model_key in _token_stats:
                                    _token_stats[model_key][0] += counts[0]
                                    _token_stats[model_key][1] += counts[1]
                                else:
                                    _token_stats[model_key] = list(counts)
                        _save_token_stats()
                        summary["token_stats_merged"] = True
                    except Exception as exc:
                        summary["errors"].append(f"Token stats merge error: {exc}")
        finally:
            Path(tmp.name).unlink(missing_ok=True)

        return summary

    result = await asyncio.to_thread(_do_import)
    has_errors = bool(result.get("errors"))
    return JSONResponse(result, status_code=207 if has_errors else 200)


if __name__ == "__main__":
    import socket
    import uvicorn

    host = HOST
    port = PORT

    lan_ip = None
    if host in {"0.0.0.0", "::"}:
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            lan_ip = s.getsockname()[0]
            s.close()
        except Exception:
            lan_ip = "127.0.0.1"

    log.info("══════════════════════════════════════════")
    log.info("  OfflineAI")
    log.info(f"  Local:    http://127.0.0.1:{port}")
    if lan_ip:
        log.info(f"  Network:  http://{lan_ip}:{port}")
        if AUTH_REQUIRED:
            log.info(f"  Token:    {AUTH_TOKEN}")
    else:
        log.info("  Network:  disabled (set OFFLINEAI_HOST=0.0.0.0 to expose)")
    log.info("══════════════════════════════════════════")
    log.info("Make sure Ollama is running: ollama serve")
    log.info(f"Make sure model is available: ollama pull {FALLBACK_MODEL}")
    uvicorn.run(app, host=host, port=port)
